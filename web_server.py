import asyncio
import csv
import gzip
import hashlib
import json
import mmap
import os
import pickle
import re
import socket
import subprocess
import threading
import time
import unicodedata
from collections import defaultdict, OrderedDict
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from contextlib import asynccontextmanager
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

from dotenv import load_dotenv

SYSTEM_VERSION = "1.1.8"
# File Version: 1.8.0
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, field_validator
import io
try:
    from colorama import Fore, Style, init as colorama_init
except Exception:
    Fore = None
    Style = None
    colorama_init = None

# --- 外部依存ロジック ---
from docx import Document
from openpyxl import load_workbook
from pdfminer.high_level import extract_pages, extract_text as miner_extract_text
from pdfminer.layout import LAParams, LTTextBoxHorizontal, LTTextBoxVertical, LTTextContainer
from pdfminer.pdfpage import PDFPage

# --- 定数 ---
BASE_DIR = Path(__file__).resolve().parent
INDEX_DIR = BASE_DIR / "indexes"
INDEX_DIR.mkdir(exist_ok=True)

ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls"}
INDEX_VERSION = "v3"
SEARCH_CACHE_VERSION = "v2"
CACHE_DIR = BASE_DIR / "cache"
CACHE_DIR.mkdir(exist_ok=True)
FIXED_CACHE_INDEX = CACHE_DIR / "fixed_cache_index.json"
QUERY_STATS_PATH = CACHE_DIR / "query_stats.json"

# --- 世代ディレクトリ定数 ---
BUILD_DIR = INDEX_DIR / ".build"
BUILD_DIR.mkdir(exist_ok=True)
CURRENT_POINTER_FILE = INDEX_DIR / "current.txt"

# --- 段階的ハッシング定数 ---
FILE_STATE_FILENAME = "file_state.jsonl"

# --- 検索表示定数 ---
SNIPPET_PREFIX_CHARS = 40
SNIPPET_TOTAL_LENGTH = 160
DETAIL_CONTEXT_PREFIX = 500
DETAIL_WINDOW_SIZE = 2000

# --- 検索パフォーマンス定数 ---
SEARCH_ENTRIES_CHUNK_THRESHOLD = 2000
HEARTBEAT_TTL_DEFAULT_SEC = 90

# --- キャッシュデフォルト定数 ---
DEFAULT_CACHE_MAX_ENTRIES = 200
DEFAULT_CACHE_MAX_MB = 200
DEFAULT_CACHE_MAX_RESULT_KB = 2000


# --- ユーティリティ ---
CID_PATTERN = re.compile(r"\(cid:\d+\)", re.IGNORECASE)
INVISIBLE_SEPARATORS_PATTERN = re.compile(r"[\u200b\u200c\u200d\u2060\ufeff\u00a0\u202f]")
JP_CHAR_CLASS = r"\u3040-\u30FF\u3400-\u4DBF\u4E00-\u9FFF\uF900-\uFAFF\uFF66-\uFF9F"
OCR_PAGE_PATTERN = re.compile(r"^\s*-{2,}\s*Page\s*(\d+)?\s*-{2,}\s*$", re.IGNORECASE)


def strip_cid(text: str) -> str:
    if not text:
        return ""
    return CID_PATTERN.sub("", text)


def normalize_invisible_separators(text: str) -> str:
    if not text:
        return ""
    return INVISIBLE_SEPARATORS_PATTERN.sub(" ", text)


def build_flexible_keyword_regex(keyword: str) -> re.Pattern | None:
    compact = re.sub(r"\s+", "", keyword or "")
    if len(compact) < 2 or len(compact) > 64:
        return None
    parts = [re.escape(ch) for ch in compact]
    pattern = r"\s*".join(parts)
    return re.compile(pattern, re.IGNORECASE)


def apply_space_mode(text: str, mode: str) -> str:
    if not text:
        return ""
    if mode == "all":
        return re.sub(r"\s+", "", text)
    if mode == "jp":
        pattern = rf"(?<=[{JP_CHAR_CLASS}])\s+(?=[{JP_CHAR_CLASS}])"
        return re.sub(pattern, "", text)
    return text


def normalize_text(text: str) -> str:
    """キーワード検索のための正規化."""
    if not text:
        return ""
    text = strip_cid(text)
    text = unicodedata.normalize("NFKC", text)
    text = normalize_invisible_separators(text)
    text = text.lower()
    text = re.sub(r"[\n\t\r]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_text_minimal(text: str) -> str:
    """改行/不可視の最小整形のみ行う."""
    if not text:
        return ""
    text = strip_cid(text)
    text = normalize_invisible_separators(text)
    text = re.sub(r"[\n\t\r]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_text_nfkc_casefold(text: str, compress_spaces: bool = True) -> str:
    """NFKC + casefold の正規化."""
    if not text:
        return ""
    text = strip_cid(text)
    text = unicodedata.normalize("NFKC", text)
    text = normalize_invisible_separators(text)
    text = text.casefold()
    text = re.sub(r"[\n\t\r]", " ", text)
    if compress_spaces:
        text = re.sub(r"\s+", " ", text)
    return text


def encode_norm_text(text: str) -> bytes:
    if not text:
        return b""
    return text.encode("utf-8")


def normalize_snippet_text(text: str) -> str:
    text = strip_cid(text)
    text = normalize_invisible_separators(text)
    text = text.replace("\r", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_detail_text(text: str) -> str:
    text = strip_cid(text)
    text = normalize_invisible_separators(text)
    text = text.replace("\r", "")
    text = re.sub(r"\n(?!　)", " ", text)
    text = re.sub(r"[ \t]+", " ", text).strip()
    return text


def build_detail_text(raw_text: str, raw_hit_pos: int) -> str:
    detail_pos = raw_hit_pos if raw_hit_pos != -1 else 0
    detail_start = max(0, detail_pos - DETAIL_CONTEXT_PREFIX)
    detail_end = min(len(raw_text), detail_start + DETAIL_WINDOW_SIZE)
    detail_text_raw = raw_text[detail_start:detail_end]
    return normalize_detail_text(detail_text_raw)


def get_ipv4_addresses() -> List[str]:
    ips: set[str] = set()
    try:
        hostname = socket.gethostname()
        for item in socket.getaddrinfo(hostname, None, socket.AF_INET):
            ip = item[4][0]
            if ip and not ip.startswith("127."):
                ips.add(ip)
        for ip in socket.gethostbyname_ex(hostname)[2]:
            if ip and not ip.startswith("127."):
                ips.add(ip)
    except Exception:
        pass

    if os.name == "nt":
        try:
            output = subprocess.check_output(["ipconfig"], text=True, errors="ignore")
            for line in output.splitlines():
                line = line.strip()
                if "IPv4 Address" in line or "IPv4 アドレス" in line:
                    parts = line.split(":")
                    if len(parts) >= 2:
                        ip = parts[-1].strip()
                        if ip and not ip.startswith("127."):
                            ips.add(ip)
        except Exception:
            pass
    else:
        try:
            output = subprocess.check_output(["/sbin/ip", "-4", "addr"], text=True, errors="ignore")
            for match in re.finditer(r"\\binet\\s+(\\d+\\.\\d+\\.\\d+\\.\\d+)", output):
                ip = match.group(1)
                if ip and not ip.startswith("127."):
                    ips.add(ip)
        except Exception:
            pass
    return sorted(ips)


def split_by_ocr_markers(text: str) -> Dict[int, str]:
    lines = text.split("\n")
    pages: Dict[int, str] = {}
    buffer: List[str] = []
    current_page = 1
    found_marker = False
    for line in lines:
        match = OCR_PAGE_PATTERN.match(line)
        if match:
            found_marker = True
            if buffer:
                pages[current_page] = "\n".join(buffer).strip()
                buffer = []
            page_hint = match.group(1)
            if page_hint and page_hint.isdigit():
                current_page = int(page_hint)
            else:
                current_page += 1
            continue
        buffer.append(line)
    if buffer:
        pages[current_page] = "\n".join(buffer).strip()
    if not found_marker:
        return {}
    return {k: v for k, v in pages.items() if v}


def find_break_pos(text: str, start: int, target: int, min_size: int, max_size: int) -> int:
    end_limit = min(len(text), start + max_size)
    target_pos = min(len(text), start + target)
    window_start = min(len(text), start + min_size)
    if window_start >= end_limit:
        return end_limit

    primary = set(["\n", " ", "\t"])
    secondary = set("。．、，.!?？！)]}」』】）〕〉》】○×△□◇■-―—")

    def scan_for(chars: set) -> int | None:
        for pos in range(target_pos, end_limit):
            if text[pos] in chars:
                return pos
        for pos in range(target_pos - 1, window_start - 1, -1):
            if text[pos] in chars:
                return pos
        return None

    hit = scan_for(primary)
    if hit is None:
        hit = scan_for(secondary)
    if hit is None:
        return min(start + target, len(text))
    while hit + 1 < len(text) and text[hit + 1] in primary:
        hit += 1
    return hit + 1


def split_text_by_size(
    text: str,
    target: int = 1200,
    min_size: int = 600,
    max_size: int = 1800,
    overlap: int = 200,
) -> Dict[int, str]:
    pages: Dict[int, str] = {}
    pos = 0
    page_num = 1
    length = len(text)
    if length <= max_size:
        cleaned = text.strip()
        return {1: cleaned} if cleaned else {}
    while pos < length:
        remaining = length - pos
        if remaining <= max_size:
            chunk = text[pos:].strip()
            if chunk:
                pages[page_num] = chunk
            break
        split_pos = find_break_pos(text, pos, target, min_size, max_size)
        if split_pos <= pos:
            split_pos = min(pos + target, length)
        if split_pos <= pos:
            break
        chunk = text[pos:split_pos].strip()
        if chunk:
            pages[page_num] = chunk
            page_num += 1
        if overlap > 0:
            pos = max(split_pos - overlap, pos + 1)
        else:
            pos = split_pos
    return pages


def split_non_pdf_text(text: str) -> Dict[int, str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    pages = split_by_ocr_markers(normalized)
    if pages:
        return pages
    return split_text_by_size(normalized)


def read_text_file_safe(path: str) -> str:
    encodings = ["utf-8", "cp932", "shift_jis", "euc-jp"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""


def read_csv_file_safe(path: str) -> str:
    encodings = ["utf-8", "cp932", "shift_jis"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc, newline="") as f:
                reader = csv.reader(f)
                return "\n".join([" ".join(row) for row in reader])
        except Exception:
            continue
    return ""


def read_excel_rows_xlsx(path: str) -> List[Tuple[str, int, str]]:
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return []

    rows: List[Tuple[str, int, str]] = []
    try:
        for ws in wb.worksheets:
            sheet_name = ws.title or "Sheet"
            row_idx = 0
            for row in ws.iter_rows(values_only=True):
                row_idx += 1
                cells = [
                    str(cell).strip()
                    for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if cells:
                    rows.append((sheet_name, row_idx, " ".join(cells)))
    except Exception:
        return []
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return rows


def read_excel_rows_xls(path: str) -> List[Tuple[str, int, str]]:
    try:
        import xlrd
    except Exception:
        return []
    try:
        wb = xlrd.open_workbook(path)
    except Exception:
        return []

    rows: List[Tuple[str, int, str]] = []
    try:
        for sheet in wb.sheets():
            sheet_name = sheet.name or "Sheet"
            for r in range(sheet.nrows):
                row_values = [
                    str(sheet.cell_value(r, c)).strip()
                    for c in range(sheet.ncols)
                    if str(sheet.cell_value(r, c)).strip()
                ]
                if row_values:
                    rows.append((sheet_name, r + 1, " ".join(row_values)))
    except Exception:
        return []
    return rows


def split_excel_rows(rows: List[Tuple[str, int, str]], rows_per_page: int = 40) -> Dict[str, str]:
    pages: Dict[str, str] = {}
    by_sheet: Dict[str, List[Tuple[int, str]]] = {}
    for sheet, row_idx, text in rows:
        by_sheet.setdefault(sheet, []).append((row_idx, text))

    for sheet, sheet_rows in by_sheet.items():
        sheet_rows.sort(key=lambda x: x[0])
        for i in range(0, len(sheet_rows), rows_per_page):
            chunk = sheet_rows[i : i + rows_per_page]
            start_row = chunk[0][0]
            end_row = chunk[-1][0]
            label = f"{sheet} 行 {start_row}-{end_row}"
            pages[label] = "\n".join(text for _, text in chunk).strip()
    return pages


def iter_text_containers(obj):
    if isinstance(obj, LTTextContainer):
        yield obj
        return
    for child in getattr(obj, "_objs", []):
        yield from iter_text_containers(child)


def order_text_boxes(boxes: List[Dict], page_bbox: Tuple[float, float, float, float]) -> str:
    if not boxes:
        return ""
    x0, y0, x1, y1 = page_bbox
    page_mid = (x0 + x1) / 2.0
    vertical_count = sum(1 for b in boxes if b["vertical"])
    horizontal_count = len(boxes) - vertical_count
    use_vertical = vertical_count > horizontal_count and vertical_count >= 3

    if use_vertical:
        ordered = sorted(boxes, key=lambda b: (-b["x0"], -b["y1"]))
    else:
        left = [b for b in boxes if b["cx"] < page_mid]
        right = [b for b in boxes if b["cx"] >= page_mid]
        is_two_column = (
            len(boxes) >= 6
            and len(left) >= 2
            and len(right) >= 2
            and 0.3 < (len(left) / len(boxes)) < 0.7
        )
        if is_two_column:
            left.sort(key=lambda b: (-b["y1"], b["x0"]))
            right.sort(key=lambda b: (-b["y1"], b["x0"]))
            ordered = left + right
        else:
            ordered = sorted(boxes, key=lambda b: (-b["y1"], b["x0"]))

    return "\n".join(b["text"].strip() for b in ordered if b["text"].strip())


def extract_pdf_text_by_layout(file_path: str) -> Dict[int, str]:
    page_texts: Dict[int, str] = {}
    laparams = LAParams()
    for page_num, layout in enumerate(extract_pages(file_path, laparams=laparams), start=1):
        boxes = []
        for container in iter_text_containers(layout):
            text = container.get_text()
            if not text or not text.strip():
                continue
            x0, y0, x1, y1 = container.bbox
            boxes.append(
                {
                    "text": text,
                    "x0": x0,
                    "y0": y0,
                    "x1": x1,
                    "y1": y1,
                    "cx": (x0 + x1) / 2.0,
                    "vertical": isinstance(container, LTTextBoxVertical),
                }
            )
        if boxes:
            ordered_text = order_text_boxes(boxes, layout.bbox)
            if ordered_text:
                page_texts[page_num] = ordered_text
    return page_texts


def extract_text_from_file(file_path: str) -> Dict[int, str]:
    """各ファイルからテキストを抽出（PDFは高精度モード固定）."""
    texts, _ = extract_text_from_file_with_reason(file_path)
    return texts


def empty_reason(file_path: str, fallback: str) -> str:
    try:
        if os.path.getsize(file_path) == 0:
            return "空ファイル"
    except Exception:
        return fallback
    return fallback


def extract_text_from_file_with_reason(file_path: str) -> Tuple[Dict[int, str], str]:
    """抽出結果と理由を返す."""
    path_obj = Path(file_path)
    ext = path_obj.suffix.lower()
    page_texts: Dict[int, str] = {}
    file_path_str = str(file_path)
    reason = ""

    try:
        if ext == ".pdf":
            try:
                page_texts = extract_pdf_text_by_layout(file_path_str)
            except Exception:
                page_texts = {}
            if not page_texts:
                laparams = LAParams()
                with open(file_path_str, "rb") as fp:
                    pages = list(PDFPage.get_pages(fp, check_extractable=False))
                    total_pages = len(pages)

                with open(file_path_str, "rb") as fp:
                    for i in range(total_pages):
                        try:
                            text = miner_extract_text(fp, page_numbers=[i], laparams=laparams)
                            if text and text.strip():
                                page_texts[i + 1] = text
                        except Exception:
                            continue
            if not page_texts:
                reason = empty_reason(file_path_str, "PDF抽出失敗")

        elif ext == ".docx":
            doc = Document(file_path_str)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            content = "\n".join(full_text)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "内容が空")

        elif ext == ".txt":
            content = read_text_file_safe(file_path_str)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")

        elif ext == ".csv":
            content = read_csv_file_safe(file_path_str)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
        elif ext == ".xlsx":
            rows = read_excel_rows_xlsx(file_path_str)
            if rows:
                page_texts = split_excel_rows(rows)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
        elif ext == ".xls":
            rows = read_excel_rows_xls(file_path_str)
            if rows:
                page_texts = split_excel_rows(rows)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
    except Exception as exc:
        return {}, f"例外:{type(exc).__name__}"

    if page_texts:
        return page_texts, ""
    return {}, reason or "抽出結果が空"


def _find_raw_hit_position(
    raw_for_positions: str,
    raw_keywords: List[str],
    space_mode: str,
) -> int:
    """生テキスト内でキーワードの位置を検索する共通ヘルパー."""
    raw_hit_pos = -1
    raw_lower = raw_for_positions.lower()
    for kw in raw_keywords:
        if not kw:
            continue
        raw_idx = raw_lower.find(kw.lower())
        if raw_idx != -1:
            raw_hit_pos = raw_idx
            break

    if raw_hit_pos == -1 and space_mode != "none":
        for kw in raw_keywords:
            if not kw:
                continue
            flex = build_flexible_keyword_regex(kw)
            if not flex:
                continue
            match = flex.search(raw_for_positions)
            if match:
                raw_hit_pos = match.start()
                break

    return raw_hit_pos


def _build_search_result(
    entry: Dict,
    raw_text: str,
    raw_for_positions: str,
    raw_hit_pos: int,
    include_detail: bool,
) -> Dict:
    """検索結果のスニペットと詳細を生成する共通ヘルパー."""
    snippet_start = max(0, raw_hit_pos - SNIPPET_PREFIX_CHARS) if raw_hit_pos != -1 else 0
    snippet_end = min(len(raw_for_positions), snippet_start + SNIPPET_TOTAL_LENGTH)
    snippet_raw = raw_for_positions[snippet_start:snippet_end]
    snippet = f"...{normalize_snippet_text(snippet_raw)}..."

    file_id = entry.get("fileId") or file_id_from_path(entry["path"])
    page_raw = entry.get("pageRaw", entry.get("page"))
    hit_pos = raw_hit_pos if raw_hit_pos >= 0 else 0
    result = {
        "file": entry["file"],
        "path": entry["path"],
        "displayPath": entry.get(
            "displayPath", display_path_for_path(entry["path"], host_aliases)
        ),
        "page": entry["page"],
        "context": snippet,
        "detail_key": {
            "file_id": file_id,
            "page": page_raw,
            "hit_pos": hit_pos,
        },
        "folderId": entry["folderId"],
        "folderName": entry["folderName"],
    }
    if include_detail:
        result["detail"] = build_detail_text(raw_text, hit_pos)
    return result


def search_text_logic(
    entry: Dict,
    norm_keyword_groups: List[List[str]],
    raw_keywords: List[str],
    search_mode: str,
    range_limit: int,
    space_mode: str,
    normalize_mode: str,
    include_detail: bool,
) -> List[Dict]:
    if not norm_keyword_groups:
        return []

    raw_text = entry["raw"]
    raw_for_positions = normalize_invisible_separators(raw_text)
    if normalize_mode == "normalized":
        if space_mode == "all":
            norm_text = entry.get("norm_cf_all") or entry["norm_all"]
        elif space_mode == "jp":
            norm_text = entry.get("norm_cf_jp") or entry["norm_jp"]
        else:
            norm_text = entry.get("norm_cf") or entry["norm"]
    else:
        if space_mode == "all":
            norm_text = entry.get("norm_strict_all") or entry["norm_all"]
        elif space_mode == "jp":
            norm_text = entry.get("norm_strict_jp") or entry["norm_jp"]
        else:
            norm_text = entry.get("norm_strict") or entry["norm"]
    is_match = False
    match_pos = -1

    if search_mode == "OR":
        for group in norm_keyword_groups:
            for k in group:
                if not k:
                    continue
                idx = norm_text.find(k)
                if idx != -1:
                    is_match = True
                    match_pos = idx
                    break
            if is_match:
                break
    elif search_mode == "AND":
        for group in norm_keyword_groups:
            if not any(k in norm_text for k in group if k):
                return []
        if range_limit == 0 or len(norm_keyword_groups) == 1:
            first_group = norm_keyword_groups[0]
            first_kw = next((k for k in first_group if k), "")
            is_match = True
            match_pos = norm_text.find(first_kw) if first_kw else 0
        else:
            positions = []
            for group_id, group in enumerate(norm_keyword_groups):
                for kw in group:
                    if not kw:
                        continue
                    start = 0
                    while True:
                        idx = norm_text.find(kw, start)
                        if idx == -1:
                            break
                        positions.append((idx, group_id))
                        start = idx + 1
            positions.sort()
            left = 0
            count_map = defaultdict(int)
            unique_count = 0
            target_unique = len(norm_keyword_groups)
            min_window_len = float("inf")
            best_start_pos = -1
            for right in range(len(positions)):
                r_idx, r_group_id = positions[right]
                if count_map[r_group_id] == 0:
                    unique_count += 1
                count_map[r_group_id] += 1
                while unique_count == target_unique:
                    l_idx, l_group_id = positions[left]
                    current_len = r_idx - l_idx
                    if current_len <= range_limit:
                        is_match = True
                        if current_len < min_window_len:
                            min_window_len = current_len
                            best_start_pos = l_idx
                    count_map[l_group_id] -= 1
                    if count_map[l_group_id] == 0:
                        unique_count -= 1
                    left += 1
            if is_match and best_start_pos != -1:
                match_pos = best_start_pos
        if not is_match:
            return []

    if not is_match:
        return []

    raw_hit_pos = _find_raw_hit_position(raw_for_positions, raw_keywords, space_mode)
    return [_build_search_result(entry, raw_text, raw_for_positions, raw_hit_pos, include_detail)]


def search_entries_chunk(
    entries: List[Dict],
    norm_keyword_groups: List[List[str]],
    raw_keywords: List[str],
    search_mode: str,
    range_limit: int,
    space_mode: str,
    normalize_mode: str,
    include_detail: bool,
) -> List[Dict]:
    chunk_results: List[Dict] = []
    for entry in entries:
        hit = search_text_logic(
            entry,
            norm_keyword_groups,
            raw_keywords,
            search_mode,
            range_limit,
            space_mode,
            normalize_mode,
            include_detail,
        )
        if hit:
            chunk_results.extend(hit)
    return chunk_results


def build_query_groups(query: str, space_mode: str, normalize_mode: str) -> Tuple[List[str], List[List[str]]]:
    keywords = [k for k in (query or "").split() if k.strip()]
    normalizer = normalize_text_nfkc_casefold if normalize_mode == "normalized" else normalize_text_minimal
    norm_keyword_groups = [
        [apply_space_mode(normalizer(k).strip(), space_mode)]
        for k in keywords
        if k.strip()
    ]
    return keywords, norm_keyword_groups




def check_range_match(norm_text: str, norm_keyword_groups: List[List[str]], range_limit: int) -> bool:
    positions = []
    for group_id, group in enumerate(norm_keyword_groups):
        for kw in group:
            if not kw:
                continue
            start = 0
            while True:
                idx = norm_text.find(kw, start)
                if idx == -1:
                    break
                positions.append((idx, group_id))
                start = idx + 1
    positions.sort()
    left = 0
    count_map = defaultdict(int)
    unique_count = 0
    target_unique = len(norm_keyword_groups)
    for right in range(len(positions)):
        r_idx, r_group_id = positions[right]
        if count_map[r_group_id] == 0:
            unique_count += 1
        count_map[r_group_id] += 1
        while unique_count == target_unique:
            l_idx, l_group_id = positions[left]
            current_len = r_idx - l_idx
            if current_len <= range_limit:
                return True
            count_map[l_group_id] -= 1
            if count_map[l_group_id] == 0:
                unique_count -= 1
            left += 1
    return False


def search_text_logic_shared(
    entry: Dict,
    mm: mmap.mmap,
    norm_keyword_groups: List[List[str]],
    norm_keyword_groups_bytes: List[List[bytes]],
    raw_keywords: List[str],
    search_mode: str,
    range_limit: int,
    space_mode: str,
    normalize_mode: str,
    include_detail: bool,
) -> List[Dict]:
    if not norm_keyword_groups_bytes:
        return []

    if normalize_mode == "normalized":
        if space_mode == "all":
            norm_start = entry.get("norm_cf_all_off", entry["norm_all_off"])
            norm_len = entry.get("norm_cf_all_len", entry["norm_all_len"])
        elif space_mode == "jp":
            norm_start = entry.get("norm_cf_jp_off", entry["norm_jp_off"])
            norm_len = entry.get("norm_cf_jp_len", entry["norm_jp_len"])
        else:
            norm_start = entry.get("norm_cf_off", entry["norm_off"])
            norm_len = entry.get("norm_cf_len", entry["norm_len"])
    else:
        if space_mode == "all":
            norm_start = entry.get("norm_strict_all_off", entry["norm_all_off"])
            norm_len = entry.get("norm_strict_all_len", entry["norm_all_len"])
        elif space_mode == "jp":
            norm_start = entry.get("norm_strict_jp_off", entry["norm_jp_off"])
            norm_len = entry.get("norm_strict_jp_len", entry["norm_jp_len"])
        else:
            norm_start = entry.get("norm_strict_off", entry["norm_off"])
            norm_len = entry.get("norm_strict_len", entry["norm_len"])
    norm_end = norm_start + norm_len

    is_match = False

    if search_mode == "OR":
        for group in norm_keyword_groups_bytes:
            for k in group:
                if not k:
                    continue
                if mm.find(k, norm_start, norm_end) != -1:
                    is_match = True
                    break
            if is_match:
                break
    elif search_mode == "AND":
        for group in norm_keyword_groups_bytes:
            if not any(mm.find(k, norm_start, norm_end) != -1 for k in group if k):
                return []
        if range_limit == 0 or len(norm_keyword_groups_bytes) == 1:
            is_match = True
        else:
            norm_text = mm[norm_start:norm_end].decode("utf-8", errors="ignore")
            is_match = check_range_match(norm_text, norm_keyword_groups, range_limit)
        if not is_match:
            return []

    if not is_match:
        return []

    raw_start = entry["raw_off"]
    raw_len = entry["raw_len"]
    raw_end = raw_start + raw_len
    raw_bytes = mm[raw_start:raw_end]
    raw_text = raw_bytes.decode("utf-8", errors="ignore")
    raw_for_positions = normalize_invisible_separators(raw_text)

    raw_hit_pos = _find_raw_hit_position(raw_for_positions, raw_keywords, space_mode)
    return [_build_search_result(entry, raw_text, raw_for_positions, raw_hit_pos, include_detail)]


def search_entries_chunk_shared(
    entries: List[Dict],
    mm: mmap.mmap,
    norm_keyword_groups: List[List[str]],
    norm_keyword_groups_bytes: List[List[bytes]],
    raw_keywords: List[str],
    search_mode: str,
    range_limit: int,
    space_mode: str,
    normalize_mode: str,
    include_detail: bool,
) -> List[Dict]:
    chunk_results: List[Dict] = []
    for entry in entries:
        hit = search_text_logic_shared(
            entry,
            mm,
            norm_keyword_groups,
            norm_keyword_groups_bytes,
            raw_keywords,
            search_mode,
            range_limit,
            space_mode,
            normalize_mode,
            include_detail,
        )
        if hit:
            chunk_results.extend(hit)
    return chunk_results


def perform_search(
    target_ids: List[str],
    params: "SearchParams",
    norm_keyword_groups: List[List[str]],
    raw_keywords: List[str],
    worker_count: int,
    include_detail: bool,
) -> List[Dict]:
    results: List[Dict] = []

    def search_folder(fid: str) -> List[Dict]:
        folder_name = folder_states[fid]["name"]
        entries = memory_pages.get(fid)
        if entries is None:
            cache = memory_indexes.get(fid, {})
            entries = build_memory_pages(fid, folder_name, cache)
            memory_pages[fid] = entries
        if os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
            log_info(
                f"検索対象: {folder_states[fid]['path']} entries={len(entries)}"
            )
        folder_results: List[Dict] = []
        entries_iter = entries
        if len(entries_iter) >= SEARCH_ENTRIES_CHUNK_THRESHOLD:
            chunk_workers = max(1, worker_count)
            chunk_size = max(1, (len(entries_iter) + chunk_workers - 1) // chunk_workers)
            chunks = [
                entries_iter[i : i + chunk_size]
                for i in range(0, len(entries_iter), chunk_size)
            ]
            with ThreadPoolExecutor(max_workers=chunk_workers) as executor:
                futures = [
                    executor.submit(
                        search_entries_chunk,
                        chunk,
                        norm_keyword_groups,
                        raw_keywords,
                        params.mode,
                        params.range_limit,
                        params.space_mode,
                        params.normalize_mode,
                        include_detail,
                    )
                    for chunk in chunks
                ]
                for future in futures:
                    try:
                        res = future.result()
                        if res:
                            folder_results.extend(res)
                    except Exception:
                        continue
        else:
            folder_results = search_entries_chunk(
                entries_iter,
                norm_keyword_groups,
                raw_keywords,
                params.mode,
                params.range_limit,
                params.space_mode,
                params.normalize_mode,
                include_detail,
            )
        if os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
            log_info(
                f"検索結果: {folder_states[fid]['path']} hits={len(folder_results)}"
            )
        return folder_results

    for fid in target_ids:
        try:
            res = search_folder(fid)
            if res:
                results.extend(res)
        except Exception:
            continue
    return results


def init_search_worker(folder_snapshot: List[Dict], aliases: Dict[str, str]):
    global WORKER_MEMORY_PAGES, WORKER_FOLDER_META, host_aliases
    host_aliases = aliases
    pages_map: Dict[str, List[Dict]] = {}
    meta_map: Dict[str, Tuple[str, str]] = {}
    for folder in folder_snapshot:
        folder_id = folder["id"]
        folder_path = folder["path"]
        folder_name = folder["name"]
        cache = load_index_from_disk(folder_path)
        pages_map[folder_id] = build_memory_pages(folder_id, folder_name, cache)
        meta_map[folder_id] = (folder_path, folder_name)
    WORKER_MEMORY_PAGES = pages_map
    WORKER_FOLDER_META = meta_map


def init_search_worker_shared(folder_snapshot: List[Dict], aliases: Dict[str, str]):
    global WORKER_SHARED_ENTRIES, WORKER_SHARED_MM, WORKER_SHARED_FILES, host_aliases, PROCESS_SHARED
    host_aliases = aliases
    PROCESS_SHARED = True
    WORKER_SHARED_ENTRIES = {}
    WORKER_SHARED_MM = {}
    WORKER_SHARED_FILES = {}
    for folder in folder_snapshot:
        folder_id = folder["id"]
        shared_path = folder["shared_path"]
        entries = folder["entries"]
        if not entries:
            WORKER_SHARED_ENTRIES[folder_id] = []
            continue
        try:
            fh = open(shared_path, "rb")
            mm = mmap.mmap(fh.fileno(), 0, access=mmap.ACCESS_READ)
        except Exception:
            log_info(f"共有ストア読込失敗: {shared_path}")
            continue
        WORKER_SHARED_FILES[folder_id] = fh
        WORKER_SHARED_MM[folder_id] = mm
        WORKER_SHARED_ENTRIES[folder_id] = entries


def perform_search_process(
    target_ids: List[str],
    params: "SearchParams",
    norm_keyword_groups: List[List[str]],
    raw_keywords: List[str],
    worker_count: int,
    include_detail: bool,
) -> List[Dict]:
    if PROCESS_SHARED:
        return perform_search_process_shared(
            target_ids,
            params,
            norm_keyword_groups,
            raw_keywords,
            worker_count,
            include_detail,
        )

    results: List[Dict] = []

    def search_folder(fid: str) -> List[Dict]:
        entries = WORKER_MEMORY_PAGES.get(fid)
        if not entries:
            meta = WORKER_FOLDER_META.get(fid)
            if meta:
                folder_path, folder_name = meta
                cache = load_index_from_disk(folder_path)
                entries = build_memory_pages(fid, folder_name, cache)
                WORKER_MEMORY_PAGES[fid] = entries
            else:
                entries = []
        folder_results: List[Dict] = []
        entries_iter = entries
        if len(entries_iter) >= SEARCH_ENTRIES_CHUNK_THRESHOLD:
            chunk_workers = max(1, worker_count)
            chunk_size = max(1, (len(entries_iter) + chunk_workers - 1) // chunk_workers)
            chunks = [
                entries_iter[i : i + chunk_size]
                for i in range(0, len(entries_iter), chunk_size)
            ]
            with ThreadPoolExecutor(max_workers=chunk_workers) as executor:
                futures = [
                    executor.submit(
                        search_entries_chunk,
                        chunk,
                        norm_keyword_groups,
                        raw_keywords,
                        params.mode,
                        params.range_limit,
                        params.space_mode,
                        params.normalize_mode,
                        include_detail,
                    )
                    for chunk in chunks
                ]
                for future in futures:
                    try:
                        res = future.result()
                        if res:
                            folder_results.extend(res)
                    except Exception:
                        continue
        else:
            folder_results = search_entries_chunk(
                entries_iter,
                norm_keyword_groups,
                raw_keywords,
                params.mode,
                params.range_limit,
                params.space_mode,
                params.normalize_mode,
                include_detail,
            )
        return folder_results

    for fid in target_ids:
        try:
            res = search_folder(fid)
            if res:
                results.extend(res)
        except Exception:
            continue
    return results


def perform_search_process_shared(
    target_ids: List[str],
    params: "SearchParams",
    norm_keyword_groups: List[List[str]],
    raw_keywords: List[str],
    worker_count: int,
    include_detail: bool,
) -> List[Dict]:
    results: List[Dict] = []
    norm_keyword_groups_bytes = [
        [encode_norm_text(k) for k in group if k] for group in norm_keyword_groups
    ]

    def search_folder(fid: str) -> List[Dict]:
        entries = WORKER_SHARED_ENTRIES.get(fid, [])
        mm = WORKER_SHARED_MM.get(fid)
        if not mm or not entries:
            return []
        folder_results: List[Dict] = []
        entries_iter = entries
        if len(entries_iter) >= SEARCH_ENTRIES_CHUNK_THRESHOLD:
            chunk_workers = max(1, worker_count)
            chunk_size = max(1, (len(entries_iter) + chunk_workers - 1) // chunk_workers)
            chunks = [
                entries_iter[i : i + chunk_size]
                for i in range(0, len(entries_iter), chunk_size)
            ]
            with ThreadPoolExecutor(max_workers=chunk_workers) as executor:
                futures = [
                    executor.submit(
                        search_entries_chunk_shared,
                        chunk,
                        mm,
                        norm_keyword_groups,
                        norm_keyword_groups_bytes,
                        raw_keywords,
                        params.mode,
                        params.range_limit,
                        params.space_mode,
                        params.normalize_mode,
                        include_detail,
                    )
                    for chunk in chunks
                ]
                for future in futures:
                    try:
                        res = future.result()
                        if res:
                            folder_results.extend(res)
                    except Exception:
                        continue
        else:
            folder_results = search_entries_chunk_shared(
                entries_iter,
                mm,
                norm_keyword_groups,
                norm_keyword_groups_bytes,
                raw_keywords,
                params.mode,
                params.range_limit,
                params.space_mode,
                params.normalize_mode,
                include_detail,
            )
        return folder_results

    for fid in target_ids:
        try:
            res = search_folder(fid)
            if res:
                results.extend(res)
        except Exception:
            continue
    return results


def run_search_direct(
    query: str,
    params: "SearchParams",
    target_ids: List[str],
    include_detail: bool = False,
) -> Tuple[List[Dict], List[str]]:
    raw_keywords, norm_keyword_groups = build_query_groups(
        query,
        params.space_mode,
        params.normalize_mode,
    )
    worker_count = search_worker_count
    if search_execution_mode == "process" and search_executor:
        future = search_executor.submit(
            perform_search_process,
            target_ids,
            params,
            norm_keyword_groups,
            raw_keywords,
            worker_count,
            include_detail,
        )
        results = future.result()
    else:
        results = perform_search(
            target_ids,
            params,
            norm_keyword_groups,
            raw_keywords,
            worker_count,
            include_detail,
        )
    return results, raw_keywords


# --- インデックス関連 ---

# --- 世代管理ユーティリティ ---
def create_generation_uuid() -> str:
    """新しい世代UUIDを生成（タイムスタンプ + ランダム）."""
    import uuid
    timestamp = int(time.time())
    random_part = uuid.uuid4().hex[:8]
    return f"{timestamp}_{random_part}"


def get_current_generation_pointer() -> str | None:
    """currentポインターファイルから世代名を取得."""
    if not CURRENT_POINTER_FILE.exists():
        return None
    try:
        gen_name = CURRENT_POINTER_FILE.read_text(encoding="utf-8").strip()
        return gen_name if gen_name else None
    except Exception:
        return None


def set_current_generation_pointer(gen_name: str) -> None:
    """currentポインターファイルに世代名を設定（原子的）."""
    temp_file = CURRENT_POINTER_FILE.with_suffix(".tmp")
    try:
        temp_file.write_text(gen_name, encoding="utf-8")
        temp_file.replace(CURRENT_POINTER_FILE)
    except Exception as e:
        log_warn(f"currentポインター設定失敗: {e}")
        if temp_file.exists():
            temp_file.unlink()


def get_generation_dir(gen_name: str | None = None, build: bool = False) -> Path:
    """世代ディレクトリのパスを取得."""
    if gen_name is None:
        gen_name = get_current_generation_pointer()
        if gen_name is None:
            return INDEX_DIR
    base = BUILD_DIR if build else INDEX_DIR
    return base / f"gen_{gen_name}"


def get_current_generation_dir() -> Path:
    """現在の世代ディレクトリを取得（current.txt 欠落時は最新世代を自動選択）."""
    gen_name = get_current_generation_pointer()
    if gen_name:
        gen_dir = get_generation_dir(gen_name, build=False)
        if gen_dir.exists():
            return gen_dir

    # current.txt がない場合、最新の世代を自動選択
    generations = list_generations()
    if generations:
        latest_gen_name, latest_gen_dir, latest_manifest = generations[0]  # 最新（ソート済み）
        log_warn(f"current.txt が見つからないため、最新世代を使用: gen_{latest_gen_name}")
        # current.txt を復旧
        set_current_generation_pointer(latest_gen_name)
        return latest_gen_dir

    # 世代ディレクトリが一つもない場合のみ INDEX_DIR にフォールバック
    return INDEX_DIR


def create_manifest(gen_name: str, gen_dir: Path, folder_states_snapshot: Dict) -> Dict:
    """manifest.json を作成."""
    import datetime
    now = datetime.datetime.now(datetime.timezone.utc)

    folder_paths = sorted([meta["path"] for meta in folder_states_snapshot.values()])
    source_folders_hash = hashlib.sha256(
        "|".join(folder_paths).encode("utf-8")
    ).hexdigest()

    folders_info = {}
    for folder_id, meta in folder_states_snapshot.items():
        path = meta["path"]
        hashed = hashlib.sha256(path.encode("utf-8")).hexdigest()[:12]
        index_file = f"index_{hashed}_{INDEX_VERSION}.pkl.gz"
        index_path = gen_dir / index_file
        file_count = 0
        if index_path.exists():
            try:
                with gzip.open(index_path, "rb") as f:
                    data = pickle.load(f)
                    if isinstance(data, dict):
                        file_count = len(data)
            except Exception:
                pass

        folders_info[folder_id] = {
            "path": path,
            "name": meta["name"],
            "file_count": file_count,
            "index_file": index_file,
        }

    manifest = {
        "index_uuid": gen_name,
        "schema_version": INDEX_VERSION,
        "created_at": now.isoformat(),
        "created_timestamp": int(now.timestamp()),
        "source_folders_hash": source_folders_hash,
        "folders": folders_info,
    }

    return manifest


def save_manifest(gen_dir: Path, manifest: Dict) -> None:
    """manifest.json を保存."""
    manifest_path = gen_dir / "manifest.json"
    try:
        with manifest_path.open("w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
    except Exception as e:
        log_warn(f"manifest.json 保存失敗: {e}")


def load_manifest(gen_dir: Path) -> Dict | None:
    """manifest.json を読み込み."""
    manifest_path = gen_dir / "manifest.json"
    if not manifest_path.exists():
        return None
    try:
        with manifest_path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def get_current_generation_manifest() -> Dict | None:
    """現在の世代の manifest.json を取得."""
    try:
        gen_dir = get_current_generation_dir()
        if not gen_dir or not gen_dir.exists():
            return None
        return load_manifest(gen_dir)
    except Exception:
        return None


def list_generations() -> List[Tuple[str, Path, Dict | None]]:
    """すべての世代をリストアップ（世代名, パス, manifest）."""
    generations = []
    for item in INDEX_DIR.iterdir():
        if item.is_dir() and item.name.startswith("gen_"):
            gen_name = item.name[4:]  # "gen_" を除く
            manifest = load_manifest(item)
            generations.append((gen_name, item, manifest))

    # タイムスタンプでソート（新しい順）
    # manifest がない場合は gen_name からタイムスタンプを抽出してフォールバック
    def get_sort_key(item):
        gen_name, gen_dir, manifest = item
        if manifest and "created_timestamp" in manifest:
            return manifest["created_timestamp"]
        # manifest がない場合、gen_name からタイムスタンプを抽出
        # gen_name format: "timestamp_randomhex"
        try:
            timestamp_str = gen_name.split("_")[0]
            return int(timestamp_str)
        except (ValueError, IndexError):
            return 0

    generations.sort(key=get_sort_key, reverse=True)
    return generations


def cleanup_old_generations(current_gen_name: str | None, grace_sec: int = 300) -> None:
    """古い世代を削除（保持ポリシーに従う）."""
    import shutil

    keep_generations = env_int("INDEX_KEEP_GENERATIONS", 3)
    keep_days = env_int("INDEX_KEEP_DAYS", 0)
    max_bytes = env_int("INDEX_MAX_BYTES", 0)

    all_generations = list_generations()

    # 現在の世代を特定
    current_gen_dir = None
    if current_gen_name:
        current_gen_dir = get_generation_dir(current_gen_name, build=False)

    # 現在の世代を除外し、猶予期間も考慮したリストを作成
    eligible_generations = []
    for gen_name, gen_dir, manifest in all_generations:
        # 現在の世代は絶対に削除しない
        if current_gen_dir and gen_dir == current_gen_dir:
            continue

        # 猶予期間内は削除しない
        if manifest and grace_sec > 0:
            created_timestamp = manifest.get("created_timestamp", 0)
            age_sec = time.time() - created_timestamp
            if age_sec < grace_sec:
                continue

        eligible_generations.append((gen_name, gen_dir, manifest))

    # 削除候補リスト
    to_delete = []

    # 世代数チェック（新しい順にソート済みなので、keep_generations より後ろを削除）
    if keep_generations > 0 and len(eligible_generations) > keep_generations:
        for gen_name, gen_dir, manifest in eligible_generations[keep_generations:]:
            to_delete.append((gen_name, gen_dir, "世代数超過"))

    # 日数チェック
    if keep_days > 0:
        for gen_name, gen_dir, manifest in eligible_generations:
            if (gen_name, gen_dir, "世代数超過") in to_delete:
                continue  # 既に削除対象
            if manifest:
                created_timestamp = manifest.get("created_timestamp", 0)
                age_days = (time.time() - created_timestamp) / 86400
                if age_days > keep_days:
                    to_delete.append((gen_name, gen_dir, "保持期限切れ"))

    # 容量チェック（現在の世代と猶予期間中の世代を含む全体のディスク使用量を計算）
    if max_bytes > 0:
        total_bytes = 0
        keep_set = set((gen_name, gen_dir) for gen_name, gen_dir, _ in to_delete)

        # 現在の世代のサイズを計算
        if current_gen_dir and current_gen_dir.exists():
            try:
                current_bytes = sum(
                    f.stat().st_size for f in current_gen_dir.rglob("*") if f.is_file()
                )
                total_bytes += current_bytes
            except Exception:
                pass

        # すべての世代（猶予期間中も含む）のサイズを計算
        # 削除対象に既に含まれているものは除外
        for gen_name, gen_dir, manifest in all_generations:
            # 現在の世代はスキップ（既に計算済み）
            if current_gen_dir and gen_dir == current_gen_dir:
                continue
            # 削除対象のものはスキップ
            if (gen_name, gen_dir) in keep_set:
                continue
            try:
                dir_bytes = sum(
                    f.stat().st_size for f in gen_dir.rglob("*") if f.is_file()
                )
                total_bytes += dir_bytes
            except Exception:
                pass

        # 容量超過の場合、古い順に削除（猶予期間中と現在の世代は除外）
        if total_bytes > max_bytes:
            for gen_name, gen_dir, manifest in reversed(eligible_generations):
                if (gen_name, gen_dir) in keep_set:
                    continue  # 既に削除対象
                try:
                    dir_bytes = sum(
                        f.stat().st_size for f in gen_dir.rglob("*") if f.is_file()
                    )
                    total_bytes -= dir_bytes
                    to_delete.append((gen_name, gen_dir, "容量超過"))
                    if total_bytes <= max_bytes:
                        break
                except Exception:
                    pass

    # 削除実行
    for gen_name, gen_dir, reason in to_delete:
        try:
            shutil.rmtree(gen_dir)
            log_info(f"世代削除: gen_{gen_name} 理由={reason}")
        except Exception as e:
            log_warn(f"世代削除失敗: gen_{gen_name} エラー={e}")


def index_path_for(folder_path: str, gen_dir: Path | None = None) -> Path:
    """インデックスファイルのパスを取得（世代ディレクトリ対応）."""
    hashed = hashlib.sha256(folder_path.encode("utf-8")).hexdigest()[:12]
    if gen_dir is None:
        gen_dir = get_current_generation_dir()
    return gen_dir / f"index_{hashed}_{INDEX_VERSION}.pkl.gz"


def file_state_path_for(folder_path: str, gen_dir: Path | None = None) -> Path:
    """file_state.jsonl のパスを取得（世代ディレクトリ対応）."""
    hashed = hashlib.sha256(folder_path.encode("utf-8")).hexdigest()[:12]
    if gen_dir is None:
        gen_dir = get_current_generation_dir()
    return gen_dir / f"file_state_{hashed}.jsonl"


def failures_path_for(folder_path: str, gen_dir: Path | None = None) -> Path:
    """failures.json のパスを取得（世代ディレクトリ対応）."""
    hashed = hashlib.sha256(folder_path.encode("utf-8")).hexdigest()[:12]
    if gen_dir is None:
        gen_dir = get_current_generation_dir()
    return gen_dir / f"failures_{hashed}.json"


def load_file_state(folder_path: str, gen_dir: Path | None = None) -> Dict[str, Dict]:
    """file_state.jsonl をディスクから読み込み（世代ディレクトリ対応）."""
    state_path = file_state_path_for(folder_path, gen_dir)
    if not state_path.exists():
        return {}
    states = {}
    try:
        with open(state_path, "r", encoding="utf-8") as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    entry = json.loads(line)
                    # 必須キー "path" の確認
                    if "path" not in entry:
                        log_warn(f"file_state.jsonl 読み込み警告: 行{line_num} に 'path' キーがありません")
                        continue
                    states[entry["path"]] = entry
                except json.JSONDecodeError:
                    log_warn(f"file_state.jsonl 読み込み警告: 行{line_num} の JSON パースに失敗しました")
                    continue
                except Exception as e:
                    log_warn(f"file_state.jsonl 読み込み警告: 行{line_num} の処理中にエラー ({e})")
                    continue
    except Exception as e:
        log_warn(f"file_state.jsonl 読み込みエラー: {state_path} ({e})")
    return states


def save_file_state(folder_path: str, states: Dict[str, Dict], gen_dir: Path | None = None):
    """file_state.jsonl をディスクに保存（世代ディレクトリ対応）."""
    state_path = file_state_path_for(folder_path, gen_dir)
    if gen_dir is not None:
        gen_dir.mkdir(parents=True, exist_ok=True)
    try:
        with open(state_path, "w", encoding="utf-8") as f:
            for path_str in sorted(states.keys()):
                f.write(json.dumps(states[path_str], ensure_ascii=False) + "\n")
    except Exception:
        # 保存失敗時は静かにスキップ
        return


def load_failures(folder_path: str, gen_dir: Path | None = None) -> Tuple[Dict[str, str], bool]:
    """failures.json をディスクから読み込み（世代ディレクトリ対応）."""
    failures_path = failures_path_for(folder_path, gen_dir)
    if not failures_path.exists():
        return {}, False
    try:
        with open(failures_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            return {str(k): str(v) for k, v in data.items()}, True
    except Exception as e:
        log_warn(f"failures_*.json 読み込みエラー: {failures_path} ({e})")
    return {}, False


def save_failures(folder_path: str, failures: Dict[str, str], gen_dir: Path | None = None):
    """failures.json をディスクに保存（世代ディレクトリ対応）."""
    failures_path = failures_path_for(folder_path, gen_dir)
    if gen_dir is not None:
        gen_dir.mkdir(parents=True, exist_ok=True)
    temp_path = failures_path.with_suffix(".json.tmp")
    try:
        with open(temp_path, "w", encoding="utf-8") as f:
            json.dump(failures, f, ensure_ascii=False)
        os.replace(temp_path, failures_path)
    except Exception as e:
        log_warn(f"failures_*.json 保存失敗: {failures_path} ({e})")
        try:
            if temp_path.exists():
                temp_path.unlink()
        except Exception:
            pass
        return


def get_file_stat(path_str: str) -> Dict:
    """ファイルの stat 情報を取得（size, mtime_ns, inode/file_id）."""
    try:
        stat_info = os.stat(path_str)
        result = {
            "size": stat_info.st_size,
            "mtime_ns": stat_info.st_mtime_ns,
        }
        # inode (Unix系) または file_id (Windows) を追加
        if hasattr(stat_info, "st_ino") and stat_info.st_ino != 0:
            result["inode"] = stat_info.st_ino
        # Windowsの場合はfile indexを取得（可能なら）
        if os.name == "nt":
            try:
                import ctypes
                from ctypes import wintypes
                kernel32 = ctypes.windll.kernel32
                GENERIC_READ = 0x80000000
                FILE_SHARE_READ = 0x00000001
                OPEN_EXISTING = 3
                FILE_FLAG_BACKUP_SEMANTICS = 0x02000000

                handle = kernel32.CreateFileW(
                    path_str,
                    GENERIC_READ,
                    FILE_SHARE_READ,
                    None,
                    OPEN_EXISTING,
                    FILE_FLAG_BACKUP_SEMANTICS,
                    None
                )
                if handle != -1:
                    class BY_HANDLE_FILE_INFORMATION(ctypes.Structure):
                        _fields_ = [
                            ("dwFileAttributes", wintypes.DWORD),
                            ("ftCreationTime", wintypes.FILETIME),
                            ("ftLastAccessTime", wintypes.FILETIME),
                            ("ftLastWriteTime", wintypes.FILETIME),
                            ("dwVolumeSerialNumber", wintypes.DWORD),
                            ("nFileSizeHigh", wintypes.DWORD),
                            ("nFileSizeLow", wintypes.DWORD),
                            ("nNumberOfLinks", wintypes.DWORD),
                            ("nFileIndexHigh", wintypes.DWORD),
                            ("nFileIndexLow", wintypes.DWORD),
                        ]
                    info = BY_HANDLE_FILE_INFORMATION()
                    if kernel32.GetFileInformationByHandle(handle, ctypes.byref(info)):
                        file_id = (info.nFileIndexHigh << 32) | info.nFileIndexLow
                        result["file_id"] = file_id
                    kernel32.CloseHandle(handle)
            except Exception:
                pass
        return result
    except Exception:
        return {}


def compute_fast_fingerprint(path_str: str, chunk_bytes: int = 65536) -> str | None:
    """fast fingerprint を計算（先頭/末尾 N バイト + size）."""
    try:
        stat_info = os.stat(path_str)
        file_size = stat_info.st_size

        hasher = hashlib.sha256()
        hasher.update(str(file_size).encode())

        with open(path_str, "rb") as f:
            # 先頭チャンク
            head_chunk = f.read(chunk_bytes)
            hasher.update(head_chunk)

            # 末尾チャンク（ファイルが十分大きい場合）
            if file_size > chunk_bytes:
                f.seek(-min(chunk_bytes, file_size - len(head_chunk)), 2)
                tail_chunk = f.read(chunk_bytes)
                hasher.update(tail_chunk)

        return hasher.hexdigest()
    except Exception:
        return None


def compute_full_hash(path_str: str, algo: str = "sha256") -> str | None:
    """full hash を計算."""
    try:
        hasher = hashlib.new(algo)
        with open(path_str, "rb") as f:
            while chunk := f.read(1048576):  # 1MB chunks
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception:
        return None


def should_compute_full_hash(path_str: str) -> bool:
    """full hash を計算すべきかどうか判定."""
    # FULL_HASH_PATHS: カンマ区切りのパス接頭辞リスト
    full_hash_paths = os.getenv("FULL_HASH_PATHS", "").strip()
    if full_hash_paths:
        for prefix in full_hash_paths.split(","):
            prefix = prefix.strip()
            if prefix and path_str.startswith(prefix):
                return True

    # FULL_HASH_EXTS: カンマ区切りの拡張子リスト
    full_hash_exts = os.getenv("FULL_HASH_EXTS", "").strip()
    if full_hash_exts:
        ext = os.path.splitext(path_str)[1].lower()
        for target_ext in full_hash_exts.split(","):
            target_ext = target_ext.strip().lower()
            if target_ext and ext == target_ext:
                return True

    # ネットワークドライブ判定（UNCパス or SMBマウント）
    if os.name == "nt":
        # Windows: UNCパス (\\server\share\...) または ネットワークドライブ
        if path_str.startswith("\\\\") or path_str.startswith("//"):
            return True
    else:
        # Unix系: /mnt, /net, /Volumes などの典型的なマウントポイント
        mount_prefixes = ["/mnt/", "/net/", "/Volumes/"]
        for prefix in mount_prefixes:
            if path_str.startswith(prefix):
                return True

    return False


def should_reindex_file(
    path_str: str,
    current_stat: Dict,
    prev_state: Dict | None,
    diff_mode: str,
    fast_fp_bytes: int,
    full_hash_algo: str,
) -> Tuple[bool, Dict]:
    """
    段階的な差分判定を行い、再インデックスが必要かどうかを返す.

    Returns:
        (should_reindex, updated_state): 再インデックス要否と更新後の state
    """
    # 前回の状態がない場合は再インデックス必要
    if not prev_state:
        new_state = {
            "path": path_str,
            "size": current_stat.get("size"),
            "mtime_ns": current_stat.get("mtime_ns"),
            "last_indexed_at": time.time(),
            "deletion_candidate_since": None,
        }
        if "inode" in current_stat:
            new_state["inode"] = current_stat["inode"]
        if "file_id" in current_stat:
            new_state["file_id"] = current_stat["file_id"]
        return True, new_state

    # Stage 1: stat 比較
    stat_changed = False
    if current_stat.get("size") != prev_state.get("size"):
        stat_changed = True
    elif current_stat.get("mtime_ns") != prev_state.get("mtime_ns"):
        stat_changed = True
    elif "inode" in current_stat and "inode" in prev_state:
        if current_stat["inode"] != prev_state["inode"]:
            stat_changed = True
    elif "file_id" in current_stat and "file_id" in prev_state:
        if current_stat["file_id"] != prev_state["file_id"]:
            stat_changed = True

    if not stat_changed:
        # stat が変わっていない場合は、削除候補をクリアして前回の状態を保持
        updated_state = prev_state.copy()
        updated_state["deletion_candidate_since"] = None
        return False, updated_state

    # stat モードの場合はここで終了（変更あり）
    if diff_mode == "stat":
        new_state = {
            "path": path_str,
            "size": current_stat.get("size"),
            "mtime_ns": current_stat.get("mtime_ns"),
            "last_indexed_at": time.time(),
            "deletion_candidate_since": None,
        }
        if "inode" in current_stat:
            new_state["inode"] = current_stat["inode"]
        if "file_id" in current_stat:
            new_state["file_id"] = current_stat["file_id"]
        return True, new_state

    # Stage 2: fast fingerprint 比較
    if diff_mode in {"stat+fastfp", "stat+fastfp+fullhash"}:
        current_fp = compute_fast_fingerprint(path_str, fast_fp_bytes)
        prev_fp = prev_state.get("fast_fp")

        if current_fp and prev_fp and current_fp == prev_fp:
            # fast fingerprint が一致する場合は変更なし
            updated_state = prev_state.copy()
            updated_state["size"] = current_stat.get("size")
            updated_state["mtime_ns"] = current_stat.get("mtime_ns")
            if "inode" in current_stat:
                updated_state["inode"] = current_stat["inode"]
            if "file_id" in current_stat:
                updated_state["file_id"] = current_stat["file_id"]
            updated_state["deletion_candidate_since"] = None
            return False, updated_state

        # fast fingerprint が異なる、または計算できない場合
        if diff_mode == "stat+fastfp":
            new_state = {
                "path": path_str,
                "size": current_stat.get("size"),
                "mtime_ns": current_stat.get("mtime_ns"),
                "fast_fp": current_fp,
                "last_indexed_at": time.time(),
                "deletion_candidate_since": None,
            }
            if "inode" in current_stat:
                new_state["inode"] = current_stat["inode"]
            if "file_id" in current_stat:
                new_state["file_id"] = current_stat["file_id"]
            return True, new_state

    # Stage 3: full hash 比較（条件付き）
    if diff_mode == "stat+fastfp+fullhash" and should_compute_full_hash(path_str):
        current_hash = compute_full_hash(path_str, full_hash_algo)
        prev_hash = prev_state.get("full_hash")

        if current_hash and prev_hash and current_hash == prev_hash:
            # full hash が一致する場合は変更なし
            updated_state = prev_state.copy()
            updated_state["size"] = current_stat.get("size")
            updated_state["mtime_ns"] = current_stat.get("mtime_ns")
            if "inode" in current_stat:
                updated_state["inode"] = current_stat["inode"]
            if "file_id" in current_stat:
                updated_state["file_id"] = current_stat["file_id"]
            # fast fingerprint を更新
            if diff_mode in {"stat+fastfp", "stat+fastfp+fullhash"}:
                current_fp = compute_fast_fingerprint(path_str, fast_fp_bytes)
                if current_fp:
                    updated_state["fast_fp"] = current_fp
            updated_state["deletion_candidate_since"] = None
            return False, updated_state

        # full hash が異なる、または計算できない場合
        new_state = {
            "path": path_str,
            "size": current_stat.get("size"),
            "mtime_ns": current_stat.get("mtime_ns"),
            "last_indexed_at": time.time(),
            "deletion_candidate_since": None,
        }
        if "inode" in current_stat:
            new_state["inode"] = current_stat["inode"]
        if "file_id" in current_stat:
            new_state["file_id"] = current_stat["file_id"]
        if diff_mode in {"stat+fastfp", "stat+fastfp+fullhash"}:
            current_fp = compute_fast_fingerprint(path_str, fast_fp_bytes)
            if current_fp:
                new_state["fast_fp"] = current_fp
        if current_hash:
            new_state["full_hash"] = current_hash
        return True, new_state

    # それ以外の場合は変更ありとして処理
    new_state = {
        "path": path_str,
        "size": current_stat.get("size"),
        "mtime_ns": current_stat.get("mtime_ns"),
        "last_indexed_at": time.time(),
        "deletion_candidate_since": None,
    }
    if "inode" in current_stat:
        new_state["inode"] = current_stat["inode"]
    if "file_id" in current_stat:
        new_state["file_id"] = current_stat["file_id"]
    if diff_mode in {"stat+fastfp", "stat+fastfp+fullhash"}:
        current_fp = compute_fast_fingerprint(path_str, fast_fp_bytes)
        if current_fp:
            new_state["fast_fp"] = current_fp
    return True, new_state


def load_index_from_disk(folder_path: str, gen_dir: Path | None = None) -> Dict[str, Dict]:
    """インデックスをディスクから読み込み（世代ディレクトリ対応）."""
    idx_path = index_path_for(folder_path, gen_dir)
    if not idx_path.exists():
        return {}
    try:
        with gzip.open(idx_path, "rb") as f:
            data = pickle.load(f)
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_index_to_disk(folder_path: str, cache: Dict[str, Dict], gen_dir: Path | None = None):
    """インデックスをディスクに保存（世代ディレクトリ対応）."""
    idx_path = index_path_for(folder_path, gen_dir)
    # gen_dir が指定されている場合、ディレクトリを作成
    if gen_dir is not None:
        gen_dir.mkdir(parents=True, exist_ok=True)
    try:
        with gzip.open(idx_path, "wb") as f:
            pickle.dump(cache, f)
    except Exception:
        # 保存失敗時は静かにスキップ
        return


def scan_files(folder: str) -> List[Path]:
    root = Path(folder)
    files = []
    for f in root.rglob("*"):
        if f.is_file() and f.suffix.lower() in ALLOWED_EXTS:
            files.append(f)
    return files


def build_index_for_folder(folder: str, previous_failures: Dict[str, str] | None = None, gen_dir: Path | None = None, prev_gen_dir: Path | None = None, failures_loaded: bool | None = None) -> Tuple[Dict[str, Dict], Dict, Dict[str, str]]:
    """インデックス作成（差分更新 & 高精度固定）."""
    start_time = time.time()

    # 段階的ハッシング設定を読み込み
    diff_mode = os.getenv("DIFF_MODE", "stat").strip().lower()
    if diff_mode not in {"stat", "stat+fastfp", "stat+fastfp+fullhash"}:
        diff_mode = "stat"
    try:
        fast_fp_bytes = int(os.getenv("FAST_FP_BYTES", "65536").strip())
        if fast_fp_bytes <= 0:
            fast_fp_bytes = 65536
    except (ValueError, AttributeError):
        fast_fp_bytes = 65536
    full_hash_algo = os.getenv("FULL_HASH_ALGO", "sha256").strip()
    if not full_hash_algo:
        full_hash_algo = "sha256"

    # 差分更新のため、前回の世代から既存キャッシュとファイル状態を読み込む
    if prev_gen_dir is not None and prev_gen_dir.exists():
        existing_cache = load_index_from_disk(folder, prev_gen_dir)
        prev_file_states = load_file_state(folder, prev_gen_dir)
    else:
        # フォールバック: 現在の世代から読み込む（初回起動時など）
        existing_cache = load_index_from_disk(folder, gen_dir)
        prev_file_states = load_file_state(folder, gen_dir)
    failures = {k: v for k, v in (previous_failures or {}).items()}
    all_files = scan_files(folder)
    if not all_files and existing_cache:
        stats = {
            "total_files": len(existing_cache),
            "indexed_files": len(existing_cache),
            "updated_files": 0,
            "skipped_files": len(existing_cache),
        }
        elapsed = time.time() - start_time
        log_info(
            f"インデックス構築スキップ: {folder} 件数={stats['indexed_files']} "
            f"理由=ファイル一覧取得失敗 時間={elapsed:.1f}s"
        )
        failures = {k: v for k, v in failures.items() if k in existing_cache}
        return existing_cache, stats, failures
    if existing_cache and all_files:
        allow_shrink = os.getenv("REBUILD_ALLOW_SHRINK", "").strip().lower() in {"1", "true", "yes"}
        if not allow_shrink and len(all_files) < max(10, int(len(existing_cache) * 0.8)):
            stats = {
                "total_files": len(existing_cache),
                "indexed_files": len(existing_cache),
                "updated_files": 0,
                "skipped_files": len(existing_cache),
            }
            elapsed = time.time() - start_time
            log_info(
                f"インデックス構築スキップ: {folder} 件数={stats['indexed_files']} "
                f"理由=ファイル数急減 時間={elapsed:.1f}s"
            )
            failures = {k: v for k, v in failures.items() if k in existing_cache}
            return existing_cache, stats, failures

    current_map = {str(f): f for f in all_files}
    failures = {k: v for k, v in failures.items() if k in current_map}
    retry_set = set(failures.keys())
    if prev_gen_dir is not None and failures_loaded is False:
        bootstrap_set = {
            path_str
            for path_str in current_map
            if path_str in prev_file_states and path_str not in existing_cache
        }
        log_notice(
            "失敗履歴が無い/読み込めないため再試行対象を拡張: "
            f"{folder} 件数={len(bootstrap_set)} "
            f"(prev_states={len(prev_file_states)} cache={len(existing_cache)} files={len(current_map)})"
        )
        if bootstrap_set and not existing_cache and len(bootstrap_set) == len(current_map):
            log_warn(f"再試行対象が全件になります: {folder}")
        retry_set.update(bootstrap_set)

    # valid_cache は既存のインデックスから開始（削除検知のため、まだフィルタしない）
    valid_cache = dict(existing_cache)

    # 削除検知の安全策: 削除候補の処理
    current_file_states: Dict[str, Dict] = {}
    for path_str in prev_file_states:
        if path_str not in current_map:
            # ファイルが見つからない
            prev_state = prev_file_states[path_str]
            if prev_state.get("deletion_candidate_since"):
                # 2回目の不在 → 削除として扱う（インデックスから除外）
                if path_str in valid_cache:
                    del valid_cache[path_str]
                    log_warn(f"削除確定（インデックス除外）: {path_str}")
            else:
                # 1回目の不在 → 削除候補としてマーク（インデックスには残す）
                prev_state["deletion_candidate_since"] = time.time()
                current_file_states[path_str] = prev_state
                log_notice(f"削除候補: {path_str}")

    # 段階的ハッシングによる差分判定
    targets = []
    target_set = set()

    def add_target(path_str: str, path_obj: Path) -> None:
        if path_str in target_set:
            return
        targets.append(path_obj)
        target_set.add(path_str)

    for path_str, path_obj in current_map.items():
        current_stat = get_file_stat(path_str)
        if not current_stat:
            # stat 取得失敗 → 再インデックス対象
            add_target(path_str, path_obj)
            current_file_states[path_str] = {
                "path": path_str,
                "last_indexed_at": time.time(),
                "deletion_candidate_since": None,
            }
            continue

        prev_state = prev_file_states.get(path_str)
        should_reindex, new_state = should_reindex_file(
            path_str, current_stat, prev_state, diff_mode, fast_fp_bytes, full_hash_algo
        )

        if should_reindex or path_str in retry_set:
            add_target(path_str, path_obj)

        current_file_states[path_str] = new_state

    updated_data: Dict[str, Dict] = {}
    if targets:
        max_workers = max(1, os.cpu_count() or 4)
        log_info(f"インデックス更新開始: {folder} 対象={len(targets)} workers={max_workers}")
        with ProcessPoolExecutor(max_workers=max_workers) as pool:
            future_to_path = {
                pool.submit(extract_text_from_file_with_reason, str(p)): str(p) for p in targets
            }
            for future in future_to_path:
                path_str = future_to_path[future]
                try:
                    texts, reason = future.result()
                    if texts:
                        updated_data[path_str] = {
                            "mtime": os.path.getmtime(path_str),
                            "data": texts,
                            "high_acc": True,
                        }
                        if path_str in failures:
                            failures.pop(path_str, None)
                        if path_str in retry_set:
                            log_notice(f"再試行成功: {path_str}")
                    else:
                        failures[path_str] = reason or "抽出結果が空"
                        log_warn(f"インデックス失敗: {path_str} 理由={failures[path_str]}")
                except Exception:
                    failures[path_str] = "例外"
                    log_warn(f"インデックス失敗: {path_str} 理由=例外")
                    continue
        log_info(f"インデックス更新完了: {folder} 更新={len(updated_data)}")

    valid_cache.update(updated_data)
    failures = {k: v for k, v in failures.items() if k not in valid_cache}
    save_index_to_disk(folder, valid_cache, gen_dir)

    # ファイル状態を保存
    save_file_state(folder, current_file_states, gen_dir)
    save_failures(folder, failures, gen_dir)

    stats = {
        "total_files": len(all_files),
        "indexed_files": len(valid_cache),
        "updated_files": len(updated_data),
        "skipped_files": len(valid_cache) - len(updated_data),
    }
    elapsed = time.time() - start_time
    log_info(f"インデックス構築完了: {folder} 件数={stats['indexed_files']} 時間={elapsed:.1f}s")

    # 失敗情報のログ出力（運用把握のため）
    if failures:
        log_warn(f"インデックス失敗ファイル: {folder} 失敗件数={len(failures)}")
        # 詳細は SEARCH_DEBUG=1 の時のみ（大量ログ対策）
        debug_mode = os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}
        if debug_mode:
            for i, (path, reason) in enumerate(list(failures.items())[:5], 1):
                log_warn(f"  [{i}] {path}: {reason}")
            if len(failures) > 5:
                log_warn(f"  ... 他 {len(failures) - 5} 件")

    return valid_cache, stats, failures


def build_memory_pages(folder_id: str, folder_name: str, cache: Dict[str, Dict]) -> List[Dict]:
    store_normalized = env_bool("INDEX_STORE_NORMALIZED", True)
    pages: List[Dict] = []
    for path, meta in cache.items():
        file_name = os.path.basename(path)
        is_pdf = file_name.lower().endswith(".pdf")
        ext = os.path.splitext(file_name)[1].lower()
        is_excel = ext in {".xlsx", ".xls"}
        display_path = display_path_for_path(path, host_aliases)
        file_id = register_file_id(path, folder_id)
        for page_num, raw_text in meta.get("data", {}).items():
            if not raw_text:
                continue
            norm_base = normalize_text(raw_text)
            norm_strict = normalize_text_minimal(raw_text)
            norm_cf_base = (
                normalize_text_nfkc_casefold(raw_text) if store_normalized else ""
            )
            page_display = "-" if not (is_pdf or is_excel) else page_num
            entry = {
                "file": file_name,
                "path": path,
                "displayPath": display_path,
                "page": page_display,
                "pageRaw": page_num,
                "raw": raw_text,
                "norm": norm_base,
                "norm_jp": apply_space_mode(norm_base, "jp"),
                "norm_all": apply_space_mode(norm_base, "all"),
                "norm_strict": norm_strict,
                "norm_strict_jp": apply_space_mode(norm_strict, "jp"),
                "norm_strict_all": apply_space_mode(norm_strict, "all"),
                "folderId": folder_id,
                "folderName": folder_name,
                "fileId": file_id,
            }
            if store_normalized:
                entry.update(
                    {
                        "norm_cf": norm_cf_base,
                        "norm_cf_jp": apply_space_mode(norm_cf_base, "jp"),
                        "norm_cf_all": apply_space_mode(norm_cf_base, "all"),
                    }
                )
            pages.append(entry)
    return pages


def build_shared_blob(shared_path: Path, entries: List[Dict]) -> List[Dict]:
    store_normalized = env_bool("INDEX_STORE_NORMALIZED", True)
    offsets: List[Dict] = []
    total_size = 0
    for entry in entries:
        raw_b = entry["raw"].encode("utf-8")
        norm_b = encode_norm_text(entry["norm"])
        norm_jp_b = encode_norm_text(entry["norm_jp"])
        norm_all_b = encode_norm_text(entry["norm_all"])
        norm_strict_b = encode_norm_text(entry.get("norm_strict", ""))
        norm_strict_jp_b = encode_norm_text(entry.get("norm_strict_jp", ""))
        norm_strict_all_b = encode_norm_text(entry.get("norm_strict_all", ""))
        norm_cf_b = encode_norm_text(entry.get("norm_cf", "")) if store_normalized else b""
        norm_cf_jp_b = encode_norm_text(entry.get("norm_cf_jp", "")) if store_normalized else b""
        norm_cf_all_b = encode_norm_text(entry.get("norm_cf_all", "")) if store_normalized else b""
        entry_offsets = {
            "file": entry["file"],
            "path": entry["path"],
            "displayPath": entry["displayPath"],
            "page": entry["page"],
            "pageRaw": entry.get("pageRaw", entry["page"]),
            "folderId": entry["folderId"],
            "folderName": entry["folderName"],
            "fileId": entry.get("fileId") or file_id_from_path(entry["path"]),
            "raw_off": total_size,
            "raw_len": len(raw_b),
            "norm_off": total_size + len(raw_b),
            "norm_len": len(norm_b),
            "norm_jp_off": total_size + len(raw_b) + len(norm_b),
            "norm_jp_len": len(norm_jp_b),
            "norm_all_off": total_size + len(raw_b) + len(norm_b) + len(norm_jp_b),
            "norm_all_len": len(norm_all_b),
            "norm_strict_off": total_size + len(raw_b) + len(norm_b) + len(norm_jp_b) + len(norm_all_b),
            "norm_strict_len": len(norm_strict_b),
            "norm_strict_jp_off": total_size + len(raw_b) + len(norm_b) + len(norm_jp_b) + len(norm_all_b) + len(norm_strict_b),
            "norm_strict_jp_len": len(norm_strict_jp_b),
            "norm_strict_all_off": total_size + len(raw_b) + len(norm_b) + len(norm_jp_b) + len(norm_all_b) + len(norm_strict_b) + len(norm_strict_jp_b),
            "norm_strict_all_len": len(norm_strict_all_b),
        }
        extra_size = 0
        if store_normalized:
            entry_offsets.update(
                {
                    "norm_cf_off": total_size
                    + len(raw_b)
                    + len(norm_b)
                    + len(norm_jp_b)
                    + len(norm_all_b)
                    + len(norm_strict_b)
                    + len(norm_strict_jp_b)
                    + len(norm_strict_all_b),
                    "norm_cf_len": len(norm_cf_b),
                    "norm_cf_jp_off": total_size
                    + len(raw_b)
                    + len(norm_b)
                    + len(norm_jp_b)
                    + len(norm_all_b)
                    + len(norm_strict_b)
                    + len(norm_strict_jp_b)
                    + len(norm_strict_all_b)
                    + len(norm_cf_b),
                    "norm_cf_jp_len": len(norm_cf_jp_b),
                    "norm_cf_all_off": total_size
                    + len(raw_b)
                    + len(norm_b)
                    + len(norm_jp_b)
                    + len(norm_all_b)
                    + len(norm_strict_b)
                    + len(norm_strict_jp_b)
                    + len(norm_strict_all_b)
                    + len(norm_cf_b)
                    + len(norm_cf_jp_b),
                    "norm_cf_all_len": len(norm_cf_all_b),
                }
            )
            extra_size = len(norm_cf_b) + len(norm_cf_jp_b) + len(norm_cf_all_b)
        offsets.append(
            entry_offsets
        )
        total_size += (
            len(raw_b)
            + len(norm_b)
            + len(norm_jp_b)
            + len(norm_all_b)
            + len(norm_strict_b)
            + len(norm_strict_jp_b)
            + len(norm_strict_all_b)
            + extra_size
        )

    with open(shared_path, "wb") as f:
        f.truncate(total_size)
        for entry in entries:
            f.write(entry["raw"].encode("utf-8"))
            f.write(encode_norm_text(entry["norm"]))
            f.write(encode_norm_text(entry["norm_jp"]))
            f.write(encode_norm_text(entry["norm_all"]))
            f.write(encode_norm_text(entry.get("norm_strict", "")))
            f.write(encode_norm_text(entry.get("norm_strict_jp", "")))
            f.write(encode_norm_text(entry.get("norm_strict_all", "")))
            if store_normalized:
                f.write(encode_norm_text(entry.get("norm_cf", "")))
                f.write(encode_norm_text(entry.get("norm_cf_jp", "")))
                f.write(encode_norm_text(entry.get("norm_cf_all", "")))
    return offsets


def build_process_shared_store() -> List[Dict]:
    folder_snapshot: List[Dict] = []
    for folder_id, meta in folder_states.items():
        folder_path = meta["path"]
        folder_name = meta["name"]
        entries = memory_pages.get(folder_id, [])
        if not entries:
            log_info(f"共有ストア警告: 空のページ一覧 {folder_name}")
        shared_path = INDEX_DIR / f"{folder_id}_shared.bin"
        offsets = build_shared_blob(shared_path, entries)
        folder_snapshot.append(
            {
                "id": folder_id,
                "path": folder_path,
                "name": folder_name,
                "shared_path": str(shared_path),
                "entries": offsets,
            }
        )
    return folder_snapshot


# --- リクエストモデル ---
class SearchRequest(BaseModel):
    query: str = Field(..., description="検索キーワード（空白区切り）")
    mode: str = Field("AND", pattern="^(AND|OR)$")
    range_limit: int = Field(0, ge=0, le=5000)
    space_mode: str = Field("jp", pattern="^(none|jp|all)$")
    normalize_mode: str | None = Field(None)
    folders: List[str]

    @field_validator("query")
    def validate_query(cls, v: str) -> str:
        if not v.strip():
            raise ValueError("検索キーワードを指定してください")
        return v

    @field_validator("folders")
    def validate_folders(cls, v: List[str]) -> List[str]:
        if not v:
            raise ValueError("検索対象フォルダを1つ以上選択してください")
        return v

    @field_validator("normalize_mode")
    def validate_normalize_mode(cls, v: str | None) -> str | None:
        if v is None or not str(v).strip():
            return None
        mode = str(v).strip().lower()
        if mode not in {"exact", "normalized"}:
            raise ValueError("normalize_mode は exact または normalized を指定してください")
        return mode


class HeartbeatRequest(BaseModel):
    client_id: str = Field(..., min_length=8, max_length=64)

    @field_validator("client_id")
    def validate_client_id(cls, v: str) -> str:
        if not v or not v.strip():
            raise ValueError("client_id を指定してください")
        return v.strip()

# --- グローバル状態 ---
app = FastAPI(title="Preloaded Folder Search")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

STATIC_DIR = Path("static")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

load_dotenv()


def parse_host_aliases() -> Dict[str, str]:
    raw = os.getenv("SEARCH_FOLDER_ALIASES", "")
    normalized = raw.replace("\n", ";").replace("|", ";").replace(",", ";")
    aliases: Dict[str, str] = {}
    for part in normalized.split(";"):
        entry = part.strip()
        if not entry or entry.startswith("#") or "=" not in entry:
            continue
        key, value = entry.split("=", 1)
        key = key.strip().lower()
        value = value.strip()
        if key and value:
            aliases[key] = value
    return aliases


def display_path_for_path(path: str, aliases: Dict[str, str]) -> str:
    if not aliases:
        return path
    match = re.match(r"^(\\\\|//)([^\\\\/]+)([\\\\/].*)?$", path)
    if not match:
        for host, alias in aliases.items():
            pattern = re.compile(rf"(\\\\|//){re.escape(host)}(?=[\\\\/])", re.IGNORECASE)
            if pattern.search(path):
                return pattern.sub(lambda m: f"{m.group(1)}{alias}", path)
        return path
    prefix, host, rest = match.groups()
    host_key = host.lower()
    alias = aliases.get(host_key)
    if not alias:
        return path
    return f"{prefix}{alias}{rest or ''}"


def parse_configured_folders() -> List[Dict[str, str]]:
    raw = os.getenv("SEARCH_FOLDERS", "")
    normalized = raw.replace("\n", ";").replace("|", ";").replace(",", ";")
    folders: List[Dict[str, str]] = []
    for part in normalized.split(";"):
        entry = part.strip()
        if not entry:
            continue
        if "#" in entry:
            entry = entry.split("#", 1)[0].strip()
        if not entry:
            continue
        label = None
        if "=" in entry:
            label, entry = entry.split("=", 1)
            label = label.strip()
            entry = entry.strip()
        p = entry.strip().strip('"').strip("'")
        if not p:
            continue
        # smb://host/share または smb:\host\share を UNC/posix 共有パスに変換
        if p.lower().startswith("smb:"):
            without_scheme = p[4:].lstrip("/\\")
            if not without_scheme:
                continue
            if os.name == "nt":
                unc_path = without_scheme.replace("/", "\\")
                p = f"\\\\{unc_path}"
            else:
                posix_path = without_scheme.replace("\\", "/")
                p = f"//{posix_path}"
        folders.append(
            {
                "path": os.path.abspath(os.path.expanduser(p)),
                "label": label or "",
            }
        )
    return folders


host_aliases = parse_host_aliases()
configured_folders = parse_configured_folders()
folder_states: Dict[str, Dict] = {}
memory_indexes: Dict[str, Dict[str, Dict]] = {}
memory_pages: Dict[str, List[Dict]] = {}
index_failures: Dict[str, Dict[str, str]] = {}
cache_lock = threading.RLock()
memory_cache: "MemoryCache | None" = None
query_stats: Dict[str, Dict] = {}
fixed_cache_index: Dict[str, Dict] = {}
fixed_cache_dirty = False
fixed_cache_last_saved = 0.0
query_stats_dirty = False
query_stats_last_saved = 0.0
fixed_cache_rebuild_in_progress = False
fixed_cache_last_trigger = 0.0
folder_order: Dict[str, int] = {}
access_urls: List[str] = []
PROCESS_SHARED = False
WORKER_SHARED_ENTRIES: Dict[str, List[Dict]] = {}
WORKER_SHARED_MM: Dict[str, mmap.mmap] = {}
WORKER_SHARED_FILES: Dict[str, object] = {}
search_semaphore: asyncio.Semaphore | None = None
rw_lock = None
search_worker_count = 1
search_concurrency = 1
search_executor = None
search_execution_mode = "thread"
normalize_mode_warning_emitted = False
file_id_map: Dict[str, Dict[str, str]] = {}
file_id_lock = threading.RLock()
active_clients: Dict[str, float] = {}
active_clients_lock = threading.RLock()


WORKER_MEMORY_PAGES: Dict[str, List[Dict]] = {}
WORKER_FOLDER_META: Dict[str, Tuple[str, str]] = {}


@dataclass(frozen=True)
class SearchParams:
    mode: str
    range_limit: int
    space_mode: str
    normalize_mode: str


class AsyncRWLock:
    def __init__(self) -> None:
        self._readers = 0
        self._readers_lock = asyncio.Lock()
        self._writer_lock = asyncio.Lock()

    @asynccontextmanager
    async def read_lock(self):
        async with self._readers_lock:
            self._readers += 1
            if self._readers == 1:
                await self._writer_lock.acquire()
        try:
            yield
        finally:
            async with self._readers_lock:
                self._readers -= 1
                if self._readers == 0:
                    self._writer_lock.release()

    @asynccontextmanager
    async def write_lock(self):
        await self._writer_lock.acquire()
        try:
            yield
        finally:
            self._writer_lock.release()


def folder_id_from_path(path: str) -> str:
    return hashlib.sha256(path.encode("utf-8")).hexdigest()[:10]


def file_id_from_path(path: str) -> str:
    return hashlib.sha256(path.encode("utf-8")).hexdigest()[:8]


def register_file_id(path: str, folder_id: str) -> str:
    file_id = file_id_from_path(path)
    with file_id_lock:
        existing = file_id_map.get(file_id)
        if existing and existing["path"] != path:
            log_warn(f"file_id衝突: {file_id} {existing['path']} != {path}")
        else:
            file_id_map[file_id] = {"path": path, "folderId": folder_id}
    return file_id


def rebuild_file_id_map() -> None:
    with file_id_lock:
        file_id_map.clear()
        for folder_id, cache in memory_indexes.items():
            for path in cache.keys():
                file_id = file_id_from_path(path)
                existing = file_id_map.get(file_id)
                if existing and existing["path"] != path:
                    log_warn(f"file_id衝突: {file_id} {existing['path']} != {path}")
                    continue
                file_id_map[file_id] = {"path": path, "folderId": folder_id}


def resolve_page_key(data: Dict, page: str) -> object | None:
    if page is None:
        return None
    page_str = str(page)
    if page_str.isdigit():
        page_int = int(page_str)
        if page_int in data:
            return page_int
    if page_str in data:
        return page_str
    return None


def describe_folder_state() -> List[Dict]:
    items = []
    for folder_id, meta in folder_states.items():
        files = memory_indexes.get(folder_id, {})
        display_path = display_path_for_path(meta["path"], host_aliases)
        items.append(
            {
                "id": folder_id,
                "path": meta["path"],
                "displayPath": display_path,
                "name": meta["name"],
                "ready": meta.get("ready", False),
                "message": meta.get("message", ""),
                "stats": meta.get("stats", {}),
                "indexed_files": len(files),
            }
        )
    return items


def init_folder_states():
    for folder in configured_folders:
        folder_path = folder["path"]
        folder_label = folder.get("label") or ""
        folder_id = folder_id_from_path(folder_path)
        display_path = display_path_for_path(folder_path, host_aliases)
        display_name = folder_label or Path(folder_path).name or folder_path
        folder_states[folder_id] = {
            "id": folder_id,
            "path": folder_path,
            "displayPath": display_path,
            "name": display_name,
            "ready": False,
            "message": "未処理",
            "stats": {},
        }


def migrate_legacy_indexes_to_generation() -> bool:
    """既存のフラットなインデックスを初回世代に移行."""
    import shutil

    # current.txt が既に存在する場合は移行不要
    if CURRENT_POINTER_FILE.exists():
        return False

    # 既存のフラットなインデックスファイルを検索
    legacy_indexes = list(INDEX_DIR.glob(f"index_*_{INDEX_VERSION}.pkl.gz"))
    if not legacy_indexes:
        return False

    log_info(f"既存インデックスの移行を開始: {len(legacy_indexes)} ファイル")

    # 初回世代を作成
    gen_name = create_generation_uuid()
    gen_dir = get_generation_dir(gen_name, build=False)
    gen_dir.mkdir(parents=True, exist_ok=True)

    # インデックスファイルを移動
    migrated_count = 0
    for index_file in legacy_indexes:
        try:
            dest = gen_dir / index_file.name
            shutil.move(str(index_file), str(dest))
            migrated_count += 1
        except Exception as e:
            log_warn(f"インデックス移行失敗: {index_file.name} エラー={e}")

    if migrated_count > 0:
        # manifest.json を作成（簡易版）
        import datetime
        now = datetime.datetime.now(datetime.timezone.utc)
        manifest = {
            "index_uuid": gen_name,
            "schema_version": INDEX_VERSION,
            "created_at": now.isoformat(),
            "created_timestamp": int(now.timestamp()),
            "source_folders_hash": "",
            "folders": {},
            "migrated_from_legacy": True,
        }
        save_manifest(gen_dir, manifest)

        # currentポインターを設定
        set_current_generation_pointer(gen_name)
        log_info(f"既存インデックスの移行完了: {migrated_count} ファイル → gen_{gen_name}")
        return True

    return False


def build_all_indexes():
    """インデックス構築（世代ディレクトリ方式）."""
    import shutil

    log_info("インデックス構築開始")

    # 初回起動時: 既存のフラットなインデックスを世代ディレクトリに移行
    migrate_legacy_indexes_to_generation()

    # 新しい世代を作成
    gen_name = create_generation_uuid()
    build_gen_dir = get_generation_dir(gen_name, build=True)
    build_gen_dir.mkdir(parents=True, exist_ok=True)

    log_info(f"世代ディレクトリ構築: gen_{gen_name}")

    # エラー時にビルドディレクトリを確実に削除するため try-finally を使用
    try:
        # 各フォルダのインデックスを構築
        temp_indexes = {}
        temp_failures = {}

        for folder_id, meta in folder_states.items():
            path = meta["path"]
            if not os.path.isdir(path):
                meta["ready"] = False
                meta["message"] = f"フォルダが見つかりません: {path}"
                continue

            meta["message"] = "インデックス構築中..."
            log_info(f"フォルダ処理開始: {path}")

            # 差分更新のため、前回の世代ディレクトリを取得
            prev_gen_dir = get_current_generation_dir()

            prev_failures = index_failures.get(folder_id, {})
            if prev_gen_dir:
                prev_failures_disk, failures_loaded = load_failures(path, prev_gen_dir)
            else:
                prev_failures_disk, failures_loaded = {}, None
            merged_failures = dict(prev_failures_disk)
            merged_failures.update(prev_failures)
            cache, stats, failures = build_index_for_folder(
                path, merged_failures, build_gen_dir, prev_gen_dir, failures_loaded
            )

            temp_indexes[folder_id] = cache
            temp_failures[folder_id] = failures

            if os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
                log_info(f"インデックス構築: {path} files={len(cache)}")

            meta["stats"] = stats
            log_info(f"フォルダ処理完了: {path}")

        # manifest.json を作成
        manifest = create_manifest(gen_name, build_gen_dir, folder_states)
        save_manifest(build_gen_dir, manifest)
        log_info(f"manifest.json 作成完了: gen_{gen_name}")

        # ビルドディレクトリを本番ディレクトリに移動（原子的）
        final_gen_dir = get_generation_dir(gen_name, build=False)
        build_gen_dir.replace(final_gen_dir)
        log_info(f"世代ディレクトリ移動完了: gen_{gen_name}")

        # currentポインターを更新（原子的）
        old_gen_name = get_current_generation_pointer()
        set_current_generation_pointer(gen_name)
        log_info(f"currentポインター更新: {old_gen_name or 'なし'} → gen_{gen_name}")

        # メモリ内のインデックスとページを更新
        for folder_id, cache in temp_indexes.items():
            memory_indexes[folder_id] = cache
            memory_pages[folder_id] = build_memory_pages(
                folder_id, folder_states[folder_id]["name"], cache
            )
            index_failures[folder_id] = temp_failures[folder_id]
            folder_states[folder_id]["ready"] = True
            folder_states[folder_id]["message"] = "準備完了"

            if os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
                log_info(
                    f"ページ数: {folder_states[folder_id]['path']} "
                    f"files={len(cache)} pages={len(memory_pages[folder_id])}"
                )

        rebuild_file_id_map()

        # 古い世代をクリーンアップ
        grace_sec = env_int("INDEX_CLEANUP_GRACE_SEC", 300)
        cleanup_old_generations(gen_name, grace_sec)

        log_info("インデックス構築完了")

    except Exception as e:
        # エラー発生時: ビルドディレクトリを削除
        log_warn(f"インデックス構築失敗: {e}")
        if build_gen_dir.exists():
            try:
                shutil.rmtree(build_gen_dir)
                log_info(f"ビルドディレクトリ削除: {build_gen_dir}")
            except Exception as cleanup_error:
                log_warn(f"ビルドディレクトリ削除失敗: {cleanup_error}")
        raise


def install_asyncio_exception_filter():
    loop = asyncio.get_running_loop()
    original_handler = loop.get_exception_handler()

    def handler(loop, context):
        exc = context.get("exception")
        message = context.get("message", "")
        if isinstance(exc, ConnectionResetError):
            return
        if "ConnectionResetError" in message or "WinError 10054" in message:
            return
        if original_handler:
            original_handler(loop, context)
        else:
            loop.default_exception_handler(context)

    loop.set_exception_handler(handler)


def log_info(message: str):
    if colorama_init:
        colorama_init()
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")


def log_warn(message: str):
    if colorama_init:
        colorama_init()
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    if Fore and Style:
        print(f"{Fore.RED}[{timestamp}] {message}{Style.RESET_ALL}")
    else:
        print(f"[{timestamp}] {message}")


def log_notice(message: str):
    """黄色で表示する通知ログ（削除候補など、警告だが重要度は中程度）."""
    if colorama_init:
        colorama_init()
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    if Fore and Style:
        print(f"{Fore.YELLOW}[{timestamp}] {message}{Style.RESET_ALL}")
    else:
        print(f"[{timestamp}] {message}")


def log_success(message: str):
    if colorama_init:
        colorama_init()
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    if Fore and Style:
        print(f"{Fore.GREEN}[{timestamp}] {message}{Style.RESET_ALL}")
    else:
        print(f"[{timestamp}] {message}")


def colorize_url(url: str) -> str:
    if Fore and Style:
        return f"{Fore.CYAN}{url}{Style.RESET_ALL}"
    return url


def env_int(name: str, default: int) -> int:
    raw = os.getenv(name, "").strip()
    if raw.isdigit():
        return int(raw)
    return default


def env_float(name: str, default: float) -> float:
    raw = os.getenv(name, "").strip()
    try:
        return float(raw)
    except Exception:
        return default


def env_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name, "").strip().lower()
    if raw in {"1", "true", "yes"}:
        return True
    if raw in {"0", "false", "no"}:
        return False
    return default


def resolve_normalize_mode(value: str | None) -> str:
    global normalize_mode_warning_emitted
    store_normalized = env_bool("INDEX_STORE_NORMALIZED", True)
    if value:
        normalized = value.strip().lower()
        if normalized in {"exact", "normalized"}:
            if normalized == "normalized" and not store_normalized:
                if not normalize_mode_warning_emitted:
                    log_warn("正規化モードは INDEX_STORE_NORMALIZED=1 が必要なため exact にフォールバックします")
                    normalize_mode_warning_emitted = True
                return "exact"
            return normalized
    env_value = os.getenv("QUERY_NORMALIZE", "nfkc_casefold").strip().lower()
    if env_value == "nfkc_casefold":
        if store_normalized:
            return "normalized"
        if not normalize_mode_warning_emitted:
            log_warn("QUERY_NORMALIZE=nfkc_casefold には INDEX_STORE_NORMALIZED=1 が必要です")
            normalize_mode_warning_emitted = True
    return "exact"


def normalize_mode_label(mode: str) -> str:
    if mode == "normalized":
        return "ゆらぎ吸収"
    return "厳格（最小整形）"


class MemoryCache:
    def __init__(self, max_entries: int, max_bytes: int, max_result_kb: int) -> None:
        self.max_entries = max_entries
        self.max_bytes = max_bytes
        self.max_result_kb = max_result_kb
        self._data: OrderedDict[str, Dict] = OrderedDict()
        self._sizes: Dict[str, int] = {}
        self._total_bytes = 0
        self._lock = threading.RLock()

    def configure(self, max_entries: int, max_bytes: int, max_result_kb: int) -> None:
        with self._lock:
            self.max_entries = max_entries
            self.max_bytes = max_bytes
            self.max_result_kb = max_result_kb
            self._evict_if_needed()

    def _evict_if_needed(self) -> None:
        while self._data and (len(self._data) > self.max_entries or self._total_bytes > self.max_bytes):
            key, _ = self._data.popitem(last=False)
            size = self._sizes.pop(key, 0)
            self._total_bytes = max(0, self._total_bytes - size)

    def clear(self) -> None:
        with self._lock:
            self._data.clear()
            self._sizes.clear()
            self._total_bytes = 0

    def get(self, key: str) -> Dict | None:
        with self._lock:
            payload = self._data.get(key)
            if payload is None:
                return None
            self._data.move_to_end(key)
            return payload

    def set(self, key: str, payload: Dict, size_bytes: int) -> None:
        size_kb = size_bytes // 1024
        if size_kb > self.max_result_kb:
            return
        with self._lock:
            if key in self._data:
                self._data.move_to_end(key)
                prev_size = self._sizes.get(key, 0)
                self._total_bytes = max(0, self._total_bytes - prev_size)
            self._data[key] = payload
            self._sizes[key] = size_bytes
            self._total_bytes += size_bytes
            self._evict_if_needed()

    def pop(self, key: str) -> None:
        with self._lock:
            if key in self._data:
                self._data.pop(key, None)
                size = self._sizes.pop(key, 0)
                self._total_bytes = max(0, self._total_bytes - size)


def normalize_query_key(query: str) -> str:
    return " ".join((query or "").strip().split())


def cache_key_for(query: str, params: "SearchParams", target_ids: List[str]) -> str:
    payload = {
        "v": SEARCH_CACHE_VERSION,
        "q": normalize_query_key(query),
        "mode": params.mode,
        "range": params.range_limit,
        "space": params.space_mode,
        "normalize": params.normalize_mode,
        "folders": sorted(target_ids),
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def estimate_payload_bytes(payload: Dict) -> int:
    raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    return len(raw)


def normalize_folder_ids(ids: List[str]) -> List[str]:
    return sorted(ids)


def payload_matches_folders(payload: Dict, target_ids: List[str]) -> bool:
    expected = normalize_folder_ids(target_ids)
    stored = payload.get("folder_ids")
    if not isinstance(stored, list):
        return False
    return normalize_folder_ids(stored) == expected


def folder_order_map() -> Dict[str, int]:
    order: List[str] = []
    for folder in configured_folders:
        folder_path = folder.get("path")
        if not folder_path:
            continue
        order.append(folder_id_from_path(folder_path))
    return {fid: idx for idx, fid in enumerate(order)}


def apply_folder_order(payload: Dict, order_map: Dict[str, int]) -> Dict:
    results = payload.get("results")
    if not isinstance(results, list):
        return payload
    fallback = len(order_map) + 1
    results.sort(key=lambda r: order_map.get(r.get("folderId"), fallback))
    return payload


def load_fixed_cache_index() -> Dict[str, Dict]:
    if not FIXED_CACHE_INDEX.exists():
        return {}
    try:
        with open(FIXED_CACHE_INDEX, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            return data
        return {}
    except Exception:
        return {}


def save_fixed_cache_index(index: Dict[str, Dict]) -> None:
    try:
        with open(FIXED_CACHE_INDEX, "w", encoding="utf-8") as f:
            json.dump(index, f, ensure_ascii=False, indent=2)
    except Exception:
        return


def load_query_stats() -> Dict[str, Dict]:
    if not QUERY_STATS_PATH.exists():
        return {}
    try:
        with open(QUERY_STATS_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            return data
        return {}
    except Exception:
        return {}


def save_query_stats(stats: Dict[str, Dict]) -> None:
    try:
        with open(QUERY_STATS_PATH, "w", encoding="utf-8") as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
    except Exception:
        return


def touch_fixed_cache_entry(key: str) -> None:
    global fixed_cache_dirty
    entry = fixed_cache_index.get(key)
    if not entry:
        return
    entry["last_access"] = time.time()
    fixed_cache_dirty = True


def maybe_flush_fixed_cache_index(force: bool = False) -> None:
    global fixed_cache_dirty, fixed_cache_last_saved
    if not fixed_cache_dirty and not force:
        return
    now = time.time()
    if not force and now - fixed_cache_last_saved < 60:
        return
    save_fixed_cache_index(fixed_cache_index)
    fixed_cache_last_saved = now
    fixed_cache_dirty = False


def maybe_flush_query_stats(force: bool = False) -> None:
    global query_stats_dirty, query_stats_last_saved
    if not query_stats_dirty and not force:
        return
    now = time.time()
    interval = env_int("QUERY_STATS_FLUSH_SEC", 60)
    if not force and now - query_stats_last_saved < interval:
        return
    save_query_stats(query_stats)
    query_stats_last_saved = now
    query_stats_dirty = False


def fixed_cache_path(key: str, compressed: bool) -> Path:
    suffix = "json.gz" if compressed else "json"
    return CACHE_DIR / f"{key}.{suffix}"


def read_fixed_cache_payload(entry: Dict) -> Dict | None:
    path = entry.get("path")
    if not path:
        return None
    try:
        if entry.get("compressed"):
            with gzip.open(path, "rt", encoding="utf-8") as f:
                return json.load(f)
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def write_fixed_cache_payload(key: str, payload: Dict, compress: bool) -> Path | None:
    path = fixed_cache_path(key, compress)
    try:
        if compress:
            with gzip.open(path, "wt", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False)
        else:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False)
        return path
    except Exception:
        return None


def prune_cache_dir(valid_keys: set[str]) -> None:
    """
    無効なキャッシュファイルを削除する。
    メタデータファイル（fixed_cache_index.json, query_stats.json）は保護される。
    """
    # メタデータファイルを除外（削除対象外）
    metadata_files = {"fixed_cache_index", "query_stats"}
    for file in CACHE_DIR.glob("*.json*"):
        stem = file.name.split(".")[0]
        # メタデータファイルは保護
        if stem in metadata_files:
            continue
        # 有効なキャッシュキーでなければ削除
        if stem not in valid_keys:
            try:
                file.unlink()
            except Exception:
                continue


def validate_cache_integrity(cache_index: Dict[str, Dict]) -> Dict[str, Dict]:
    """
    キャッシュの整合性を検証し、不一致のエントリを削除する。
    起動時やインデックス更新後に呼び出される。
    """
    current_manifest = get_current_generation_manifest()
    if not current_manifest:
        log_warn("キャッシュ整合性検証: 現在の世代のmanifestが取得できません")
        return cache_index

    current_uuid = current_manifest.get("index_uuid")
    current_schema = current_manifest.get("schema_version")
    if not current_uuid or not current_schema:
        log_warn("キャッシュ整合性検証: manifestにindex_uuid/schema_versionがありません")
        return cache_index

    valid_cache: Dict[str, Dict] = {}
    invalid_count = 0

    for key, entry in cache_index.items():
        cached_uuid = entry.get("index_uuid")
        cached_schema = entry.get("schema_version")

        # index_uuid のチェック（欠落または不一致で無効化）
        if not cached_uuid or cached_uuid != current_uuid:
            log_info(
                f"キャッシュ削除（index_uuid不一致/欠落）: key={key[:8]}.. "
                f"cached={cached_uuid} current={current_uuid}"
            )
            invalid_count += 1
            continue

        # schema_version のチェック（欠落または不一致で無効化）
        if not cached_schema or cached_schema != current_schema:
            log_info(
                f"キャッシュ削除（schema不一致/欠落）: key={key[:8]}.. "
                f"cached={cached_schema} current={current_schema}"
            )
            invalid_count += 1
            continue

        # 有効なキャッシュとして保持
        valid_cache[key] = entry

    if invalid_count > 0:
        log_info(f"キャッシュ整合性検証完了: {invalid_count}件の古いキャッシュを削除しました")

    return valid_cache


def init_cache_settings(memory_cache: MemoryCache) -> None:
    max_entries = env_int("CACHE_MEM_MAX_ENTRIES", 200)
    max_mb = env_int("CACHE_MEM_MAX_MB", 200)
    max_result_kb = env_int("CACHE_MEM_MAX_RESULT_KB", 2000)
    memory_cache.configure(max_entries, max_mb * 1024 * 1024, max_result_kb)


def update_query_stats(
    stats: Dict[str, Dict],
    key: str,
    query: str,
    params: "SearchParams",
    target_ids: List[str],
    result_count: int,
    payload_kb: float,
    elapsed_ms: float | None,
    cached: bool,
) -> None:
    global query_stats_dirty
    now = time.time()
    entry = stats.get(key)
    if entry is None:
        entry = {
            "query": normalize_query_key(query),
            "mode": params.mode,
            "range": params.range_limit,
            "space": params.space_mode,
            "normalize": params.normalize_mode,
            "folders": sorted(target_ids),
            "count_total": 0,
            "count_uncached": 0,
            "total_time_ms": 0.0,
            "total_hits": 0,
            "total_payload_kb": 0.0,
            "last_access": now,
        }
    entry["count_total"] += 1
    entry["total_hits"] += result_count
    entry["total_payload_kb"] += payload_kb
    entry["last_access"] = now
    if not cached and elapsed_ms is not None:
        entry["count_uncached"] += 1
        entry["total_time_ms"] += elapsed_ms
    stats[key] = entry
    query_stats_dirty = True


def is_fixed_candidate(entry: Dict) -> bool:
    ttl_days = env_int("CACHE_FIXED_TTL_DAYS", 7)
    min_count = env_int("CACHE_FIXED_MIN_COUNT", 10)
    min_time = env_int("CACHE_FIXED_MIN_TIME_MS", 500)
    min_hits = env_int("CACHE_FIXED_MIN_HITS", 5000)
    min_kb = env_int("CACHE_FIXED_MIN_KB", 2000)
    now = time.time()
    if ttl_days > 0 and now - entry.get("last_access", 0) > ttl_days * 86400:
        return False
    if entry.get("count_total", 0) < min_count:
        return False
    avg_hits = entry.get("total_hits", 0) / max(1, entry.get("count_total", 1))
    avg_kb = entry.get("total_payload_kb", 0.0) / max(1, entry.get("count_total", 1))
    avg_time = entry.get("total_time_ms", 0.0) / max(1, entry.get("count_uncached", 1))
    if avg_time < min_time and avg_hits < min_hits and avg_kb < min_kb:
        return False
    return True


def rank_fixed_candidates(entries: List[Dict]) -> List[Dict]:
    def score(e: Dict) -> float:
        avg_time = e.get("total_time_ms", 0.0) / max(1, e.get("count_uncached", 1))
        avg_hits = e.get("total_hits", 0.0) / max(1, e.get("count_total", 1))
        return avg_time * max(1.0, avg_hits)
    return sorted(entries, key=score, reverse=True)


def prune_query_stats(stats: Dict[str, Dict]) -> None:
    ttl_days = env_int("QUERY_STATS_TTL_DAYS", 30)
    if ttl_days <= 0:
        return
    now = time.time()
    remove_keys = [k for k, v in stats.items() if now - v.get("last_access", 0) > ttl_days * 86400]
    for key in remove_keys:
        stats.pop(key, None)


def trigger_fixed_cache_rebuild(loop: asyncio.AbstractEventLoop, reason: str) -> None:
    global fixed_cache_rebuild_in_progress, fixed_cache_last_trigger
    cooldown = env_int("CACHE_FIXED_TRIGGER_COOLDOWN_SEC", 300)
    now = time.time()
    if fixed_cache_rebuild_in_progress:
        return
    if cooldown > 0 and now - fixed_cache_last_trigger < cooldown:
        return
    fixed_cache_rebuild_in_progress = True
    fixed_cache_last_trigger = now

    async def _run():
        nonlocal loop
        try:
            await loop.run_in_executor(None, rebuild_fixed_cache)
        finally:
            global fixed_cache_rebuild_in_progress
            fixed_cache_rebuild_in_progress = False

    log_info(f"固定キャッシュ再構築を起動します: {reason}")
    asyncio.create_task(_run())


def cpu_budget() -> int:
    env_budget = os.getenv("SEARCH_CPU_BUDGET", "").strip()
    if env_budget.isdigit():
        return max(1, int(env_budget))
    return max(1, os.cpu_count() or 4)


def total_worker_budget() -> int:
    env_budget = os.getenv("SEARCH_CPU_BUDGET", "").strip()
    if env_budget.isdigit():
        return max(1, int(env_budget))
    env_workers = os.getenv("SEARCH_WORKERS", "").strip()
    env_conc = os.getenv("SEARCH_CONCURRENCY", "").strip()
    if env_workers.isdigit() and env_conc.isdigit():
        return max(1, int(env_workers)) * max(1, int(env_conc))
    return max(1, os.cpu_count() or 4)


def heartbeat_ttl_sec() -> int:
    return max(1, env_int("HEARTBEAT_TTL_SEC", HEARTBEAT_TTL_DEFAULT_SEC))


def prune_active_clients(now: float | None = None) -> int:
    if now is None:
        now = time.time()
    ttl = heartbeat_ttl_sec()
    with active_clients_lock:
        expired = [cid for cid, ts in active_clients.items() if now - ts > ttl]
        for cid in expired:
            active_clients.pop(cid, None)
        return len(active_clients)


def record_heartbeat(client_id: str) -> int:
    now = time.time()
    with active_clients_lock:
        active_clients[client_id] = now
    return prune_active_clients(now)


def active_client_count() -> int:
    return prune_active_clients()


def init_search_settings() -> Tuple[int, int, int]:
    budget = cpu_budget()
    env_conc = os.getenv("SEARCH_CONCURRENCY", "").strip()
    concurrency = budget
    if env_conc.isdigit():
        concurrency = max(1, int(env_conc))
    env_workers = os.getenv("SEARCH_WORKERS", "").strip()
    if env_workers.isdigit():
        workers = max(1, int(env_workers))
    else:
        workers = max(1, budget // max(1, concurrency))
    return budget, concurrency, workers


def parse_rebuild_schedule(value: str) -> Tuple[str, int, int]:
    raw = (value or "").strip()
    if not raw:
        return ("", 0, 0)
    if raw.lower().endswith("h"):
        try:
            hours = int(raw[:-1])
            return ("interval", hours, 0)
        except ValueError:
            return ("", 0, 0)
    if raw.isdigit():
        return ("interval", int(raw), 0)
    try:
        parts = raw.split(":")
        if len(parts) == 2:
            hour = int(parts[0])
            minute = int(parts[1])
            if 0 <= hour <= 23 and 0 <= minute <= 59:
                return ("time", hour, minute)
    except ValueError:
        return ("", 0, 0)
    return ("", 0, 0)


async def schedule_index_rebuild():
    schedule = os.getenv("REBUILD_SCHEDULE", "").strip()
    mode, a, b = parse_rebuild_schedule(schedule)
    if not mode:
        return
    log_info(f"インデックス再構築スケジュール有効: {schedule}")
    while True:
        if mode == "interval":
            await asyncio.sleep(max(1, a) * 3600)
        elif mode == "time":
            now = time.localtime()
            target = time.mktime(
                (now.tm_year, now.tm_mon, now.tm_mday, a, b, 0, now.tm_wday, now.tm_yday, now.tm_isdst)
            )
            if target <= time.mktime(now):
                target += 86400
            await asyncio.sleep(max(1, int(target - time.mktime(now))))
        else:
            return
        if rw_lock is None:
            continue
        log_info("スケジュール再構築開始")
        async with rw_lock.write_lock():
            loop = asyncio.get_running_loop()
            await loop.run_in_executor(None, build_all_indexes)
            if search_execution_mode == "process":
                global search_executor
                if search_executor:
                    search_executor.shutdown(wait=True)
                if PROCESS_SHARED:
                    folder_snapshot = build_process_shared_store()
                    search_executor = ProcessPoolExecutor(
                        max_workers=search_concurrency,
                        initializer=init_search_worker_shared,
                        initargs=(folder_snapshot, host_aliases),
                    )
                    log_info("スケジュール再構築: process(shared) を再初期化")
                else:
                    folder_snapshot = [
                        {"id": fid, "path": meta["path"], "name": meta["name"]}
                        for fid, meta in folder_states.items()
                    ]
                    search_executor = ProcessPoolExecutor(
                        max_workers=search_concurrency,
                        initializer=init_search_worker,
                        initargs=(folder_snapshot, host_aliases),
                    )
                    log_info("スケジュール再構築: process を再初期化")
            rebuild_fixed_cache()
        log_info("スケジュール再構築完了")


def rebuild_fixed_cache():
    global fixed_cache_index

    # 現在の世代のmanifestを取得（キャッシュ整合性のため）
    current_manifest = get_current_generation_manifest()
    if not current_manifest:
        log_warn("固定キャッシュ再構築: 現在の世代のmanifestが取得できません")
        return

    current_index_uuid = current_manifest.get("index_uuid")
    current_schema_version = current_manifest.get("schema_version")
    if not current_index_uuid or not current_schema_version:
        log_warn("固定キャッシュ再構築: manifestにindex_uuid/schema_versionがありません")
        return

    log_info(f"固定キャッシュ再構築開始: index_uuid={current_index_uuid}, schema={current_schema_version}")

    prune_query_stats(query_stats)
    candidates = [v for v in query_stats.values() if is_fixed_candidate(v)]
    candidates = rank_fixed_candidates(candidates)
    max_entries = env_int("CACHE_FIXED_MAX_ENTRIES", 20)
    candidates = candidates[:max_entries]

    source_entries: List[Dict] = candidates
    if not source_entries and fixed_cache_index:
        ttl_days = env_int("CACHE_FIXED_TTL_DAYS", 7)
        now = time.time()
        source_entries = []
        for entry in fixed_cache_index.values():
            last_access = entry.get("last_access", 0)
            if ttl_days > 0 and now - last_access > ttl_days * 86400:
                continue
            source_entries.append(entry)

    new_index: Dict[str, Dict] = {}
    for entry in source_entries:
        target_ids = [fid for fid in entry.get("folders", []) if fid in folder_states and folder_states[fid].get("ready")]
        if not target_ids:
            continue
        normalize_mode = resolve_normalize_mode(entry.get("normalize"))
        params = SearchParams(entry["mode"], entry["range"], entry["space"], normalize_mode)
        try:
            results, keywords = run_search_direct(entry["query"], params, target_ids)
        except Exception:
            continue
        payload = {
            "count": len(results),
            "results": results,
            "keywords": keywords,
            "folder_ids": normalize_folder_ids(target_ids),
        }
        payload = apply_folder_order(payload, folder_order)
        size_bytes = estimate_payload_bytes(payload)
        size_kb = size_bytes / 1024
        compress = size_kb >= env_int("CACHE_COMPRESS_MIN_KB", 2000)
        key = cache_key_for(entry["query"], params, target_ids)
        path = write_fixed_cache_payload(key, payload, compress)
        if not path:
            continue
        new_index[key] = {
            "path": str(path),
            "compressed": compress,
            "payload_kb": size_kb,
            "result_count": len(results),
            "last_access": entry.get("last_access", time.time()),
            "last_built": time.time(),
            "query": entry["query"],
            "mode": entry["mode"],
            "range": entry["range"],
            "space": entry["space"],
            "normalize": normalize_mode,
            "folders": target_ids,
            "index_uuid": current_index_uuid,
            "schema_version": current_schema_version,
        }

    with cache_lock:
        fixed_cache_index = new_index
        save_fixed_cache_index(fixed_cache_index)
        if memory_cache:
            memory_cache.clear()
        global fixed_cache_dirty, fixed_cache_last_saved
        fixed_cache_dirty = False
        fixed_cache_last_saved = time.time()
        maybe_flush_query_stats(force=True)
    prune_cache_dir(set(new_index.keys()))


def per_request_workers() -> int:
    total_budget = total_worker_budget()
    max_workers = total_budget
    env_workers = os.getenv("SEARCH_WORKERS", "").strip()
    if env_workers.isdigit():
        max_workers = max(1, int(env_workers))
    clients = active_client_count()
    workers = max(1, total_budget // max(1, clients))
    if workers > max_workers:
        workers = max_workers
    return workers


@app.on_event("startup")
async def startup_event():
    init_folder_states()
    install_asyncio_exception_filter()
    global search_semaphore, rw_lock, search_worker_count, search_executor, search_execution_mode, search_concurrency, fixed_cache_index, memory_cache, query_stats, query_stats_dirty, query_stats_last_saved, folder_order
    rw_lock = AsyncRWLock()
    memory_cache = MemoryCache(
        DEFAULT_CACHE_MAX_ENTRIES,
        DEFAULT_CACHE_MAX_MB * 1024 * 1024,
        DEFAULT_CACHE_MAX_RESULT_KB,
    )
    init_cache_settings(memory_cache)
    fixed_cache_index = load_fixed_cache_index()
    # キャッシュ整合性検証（インデックス更新後の古いキャッシュを削除）
    fixed_cache_index = validate_cache_integrity(fixed_cache_index)
    save_fixed_cache_index(fixed_cache_index)
    prune_cache_dir(set(fixed_cache_index.keys()))
    query_stats = load_query_stats()
    query_stats_dirty = False
    query_stats_last_saved = time.time()
    folder_order = folder_order_map()
    budget, concurrency, workers = init_search_settings()
    search_worker_count = workers
    search_concurrency = concurrency
    search_semaphore = asyncio.Semaphore(concurrency)
    search_execution_mode = os.getenv("SEARCH_EXECUTION_MODE", "thread").strip().lower()
    loop = asyncio.get_running_loop()
    # インデックス構築はCPU負荷が高いのでスレッドで実施
    async with rw_lock.write_lock():
        await loop.run_in_executor(None, build_all_indexes)
    log_info(f"システムバージョン: {SYSTEM_VERSION} / インデックスバージョン: {INDEX_VERSION}")
    if search_execution_mode == "process":
        process_shared = os.getenv("SEARCH_PROCESS_SHARED", "1").strip().lower() not in {
            "0",
            "false",
            "no",
        }
        global PROCESS_SHARED
        PROCESS_SHARED = process_shared
        if process_shared:
            folder_snapshot = build_process_shared_store()
            search_executor = ProcessPoolExecutor(
                max_workers=concurrency,
                initializer=init_search_worker_shared,
                initargs=(folder_snapshot, host_aliases),
            )
            log_info(
                f"検索実行モード: process(shared) workers={concurrency} budget={budget}"
            )
        else:
            folder_snapshot = [
                {"id": fid, "path": meta["path"], "name": meta["name"]}
                for fid, meta in folder_states.items()
            ]
            search_executor = ProcessPoolExecutor(
                max_workers=concurrency,
                initializer=init_search_worker,
                initargs=(folder_snapshot, host_aliases),
            )
            log_info(f"検索実行モード: process workers={concurrency} budget={budget}")
    else:
        search_executor = None
        log_info(f"検索実行モード: thread concurrency={concurrency} workers={workers} budget={budget}")
    scheme = "https" if (Path(os.getenv("CERT_DIR", "certs")) / "lan-cert.pem").exists() else "http"
    urls = [f"{scheme}://{ip}:{os.getenv('PORT', '80')}" for ip in (get_ipv4_addresses() or ["0.0.0.0"])]

    async def announce_access_urls():
        await asyncio.sleep(0.2)
        for url in urls:
            log_info(f"アクセスURL: {colorize_url(url)}")

    asyncio.create_task(announce_access_urls())
    rebuild_fixed_cache()
    asyncio.create_task(schedule_index_rebuild())


@app.get("/", response_class=FileResponse)
async def read_root():
    index_file = STATIC_DIR / "index.html"
    if not index_file.exists():
        raise HTTPException(status_code=404, detail="UIが見つかりません")
    return FileResponse(index_file)


@app.get("/api/folders")
async def get_folders():
    return JSONResponse({"folders": describe_folder_state()})


@app.post("/api/heartbeat")
async def heartbeat(req: HeartbeatRequest):
    client_id = req.client_id.strip()
    if not client_id:
        raise HTTPException(status_code=400, detail="client_id が必要です")
    active_count = record_heartbeat(client_id)
    return {
        "status": "ok",
        "active_clients": active_count,
        "ttl_sec": heartbeat_ttl_sec(),
    }


@app.get("/api/folders/{folder_id}/files")
async def get_folder_files(folder_id: str, scope: str = "indexed"):
    if folder_id not in folder_states:
        raise HTTPException(status_code=404, detail="フォルダが見つかりません")
    meta = folder_states[folder_id]
    if not meta.get("ready"):
        raise HTTPException(status_code=400, detail="このフォルダは準備中です")

    base_path = meta["path"]
    cache = memory_indexes.get(folder_id, {})
    failure_map = index_failures.get(folder_id, {})
    scope = (scope or "indexed").strip().lower()
    if scope not in {"indexed", "missing", "all"}:
        raise HTTPException(status_code=400, detail="scopeが不正です")

    indexed_paths = set(cache.keys())
    files = []
    if scope == "indexed":
        for path in sorted(indexed_paths):
            rel = os.path.relpath(path, base_path)
            files.append({"path": path, "relative": rel, "indexed": True, "reason": ""})
    else:
        all_files = scan_files(base_path)
        for path_obj in sorted(all_files):
            path = str(path_obj)
            is_indexed = path in indexed_paths
            if scope == "missing" and is_indexed:
                continue
            reason = ""
            if not is_indexed:
                reason = failure_map.get(path, "理由不明")
            rel = os.path.relpath(path, base_path)
            files.append({"path": path, "relative": rel, "indexed": is_indexed, "reason": reason})
    return {"folder": folder_id, "name": meta["name"], "files": files, "scope": scope}


@app.get("/api/detail")
async def get_detail(
    file_id: str = Query(..., min_length=8, max_length=8, pattern=r"^[0-9a-f]{8}$"),
    page: str = Query(..., min_length=1),
    hit_pos: int = Query(0, ge=0),
):
    global rw_lock
    if rw_lock is None:
        raise HTTPException(status_code=503, detail="検索システムの初期化中です")
    async with rw_lock.read_lock():
        with file_id_lock:
            info = file_id_map.get(file_id)
        if not info:
            raise HTTPException(status_code=404, detail="ファイルが見つかりません")
        folder_id = info["folderId"]
        if folder_id not in folder_states or not folder_states[folder_id].get("ready"):
            raise HTTPException(status_code=400, detail="対象フォルダが準備中です")
        cache = memory_indexes.get(folder_id, {})
        meta = cache.get(info["path"])
        if not meta:
            raise HTTPException(status_code=404, detail="ファイルが見つかりません")
        data = meta.get("data", {})
        page_key = resolve_page_key(data, page)
        if page_key is None:
            raise HTTPException(status_code=404, detail="ページが見つかりません")
        raw_text = data.get(page_key)
        detail_text = build_detail_text(raw_text or "", hit_pos)
        return {"file_id": file_id, "page": page_key, "detail": detail_text}


def _try_get_memory_cache(
    cache_key: str,
    target_ids: List[str],
    req_query: str,
    params: "SearchParams",
) -> Dict | None:
    """メモリキャッシュから結果を取得。ヒットした場合は統計を更新して返す。"""
    with cache_lock:
        cached_payload = memory_cache.get(cache_key) if memory_cache else None
    if cached_payload and payload_matches_folders(cached_payload, target_ids):
        cached_payload = apply_folder_order(cached_payload, folder_order)
        cached_bytes = estimate_payload_bytes(cached_payload)
        cached_kb = cached_bytes / 1024
        update_query_stats(
            query_stats,
            cache_key,
            req_query,
            params,
            target_ids,
            cached_payload.get("count", 0),
            cached_kb,
            None,
            True,
        )
        maybe_flush_query_stats()
        return cached_payload
    elif cached_payload:
        with cache_lock:
            if memory_cache:
                memory_cache.pop(cache_key)
    return None


def _try_get_fixed_cache(
    cache_key: str,
    target_ids: List[str],
    req_query: str,
    params: "SearchParams",
) -> Dict | None:
    """固定キャッシュから結果を取得。ヒットした場合はメモリキャッシュにも保存して返す。"""
    with cache_lock:
        fixed_entry = fixed_cache_index.get(cache_key)
    if not fixed_entry:
        return None

    # キャッシュ整合性チェック: index_uuid と schema_version を検証
    current_manifest = get_current_generation_manifest()
    if current_manifest:
        current_uuid = current_manifest.get("index_uuid")
        current_schema = current_manifest.get("schema_version")
        cached_uuid = fixed_entry.get("index_uuid")
        cached_schema = fixed_entry.get("schema_version")

        # index_uuid のチェック（欠落または不一致で無効化）
        if current_uuid and (not cached_uuid or cached_uuid != current_uuid):
            log_warn(
                f"キャッシュ無効化（index_uuid不一致/欠落）: key={cache_key[:8]}.. "
                f"cached={cached_uuid} current={current_uuid}"
            )
            with cache_lock:
                fixed_cache_index.pop(cache_key, None)
                save_fixed_cache_index(fixed_cache_index)
            return None

        # schema_version のチェック（欠落または不一致で無効化）
        if current_schema and (not cached_schema or cached_schema != current_schema):
            log_warn(
                f"キャッシュ無効化（schema不一致/欠落）: key={cache_key[:8]}.. "
                f"cached={cached_schema} current={current_schema}"
            )
            with cache_lock:
                fixed_cache_index.pop(cache_key, None)
                save_fixed_cache_index(fixed_cache_index)
            return None

    payload = read_fixed_cache_payload(fixed_entry)
    if payload and payload_matches_folders(payload, target_ids):
        payload = apply_folder_order(payload, folder_order)
        cached_kb = fixed_entry.get("payload_kb")
        if cached_kb is None:
            cached_kb = estimate_payload_bytes(payload) / 1024
        update_query_stats(
            query_stats,
            cache_key,
            req_query,
            params,
            target_ids,
            payload.get("count", 0),
            cached_kb,
            None,
            True,
        )
        with cache_lock:
            if memory_cache:
                memory_cache.set(cache_key, payload, int(cached_kb * 1024))
            touch_fixed_cache_entry(cache_key)
        maybe_flush_fixed_cache_index()
        maybe_flush_query_stats()
        return payload
    with cache_lock:
        fixed_cache_index.pop(cache_key, None)
        save_fixed_cache_index(fixed_cache_index)
    return None


@app.post("/api/search")
async def search(req: SearchRequest):
    global rw_lock, search_semaphore, search_executor, search_execution_mode
    if rw_lock is None or search_semaphore is None:
        raise HTTPException(status_code=503, detail="検索システムの初期化中です")

    # Normal search (full index scan)
    async with search_semaphore:
        async with rw_lock.read_lock():
            start_time = time.time()
            available_ids = set(folder_states.keys())
            target_ids = [f for f in req.folders if f in available_ids and folder_states[f].get("ready")]
            if not target_ids:
                raise HTTPException(status_code=400, detail="有効な検索対象フォルダがありません")

            normalize_mode = resolve_normalize_mode(req.normalize_mode)
            params = SearchParams(req.mode, req.range_limit, req.space_mode, normalize_mode)
            cache_key = cache_key_for(req.query, params, target_ids)

            # キャッシュからの取得を試行
            cached_result = _try_get_memory_cache(cache_key, target_ids, req.query, params)
            if cached_result:
                return cached_result
            cached_result = _try_get_fixed_cache(cache_key, target_ids, req.query, params)
            if cached_result:
                return cached_result

            keywords, norm_keyword_groups = build_query_groups(
                req.query,
                req.space_mode,
                normalize_mode,
            )
            raw_keywords = keywords
            keywords_for_response = keywords
            worker_count = per_request_workers()
            loop = asyncio.get_running_loop()
            if os.getenv("SEARCH_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
                folder_debug = []
                for fid in target_ids:
                    name = folder_states[fid]["name"]
                    entries = memory_pages.get(fid)
                    folder_debug.append(f"{name}({fid}) entries={len(entries) if entries is not None else 'none'}")
                log_info(f"検索デバッグ: query='{req.query}' targets={folder_debug}")
            if search_execution_mode == "process":
                results = await loop.run_in_executor(
                    search_executor,
                    perform_search_process,
                    target_ids,
                    params,
                    norm_keyword_groups,
                    raw_keywords,
                    worker_count,
                    False,
                )
            else:
                results = await loop.run_in_executor(
                    None,
                    perform_search,
                    target_ids,
                    params,
                    norm_keyword_groups,
                    raw_keywords,
                    worker_count,
                    False,
                )
        elapsed = time.time() - start_time
        log_info(
            f"検索完了: folders={len(target_ids)} results={len(results)} "
            f"mode={req.mode} normalize={normalize_mode} workers={worker_count} 時間={elapsed:.2f}s"
        )

        # Get current index UUID
        current_gen_name = get_current_generation_pointer()
        index_uuid = current_gen_name or "unknown"

        payload = {
            "count": len(results),
            "results": results,
            "keywords": keywords_for_response,
            "folder_ids": normalize_folder_ids(target_ids),
            "index_uuid": index_uuid,
            "normalize_mode": normalize_mode,
        }
        payload = apply_folder_order(payload, folder_order)
        payload_bytes = estimate_payload_bytes(payload)
        payload_kb = payload_bytes / 1024
        update_query_stats(
            query_stats,
            cache_key,
            req.query,
            params,
            target_ids,
            len(results),
            payload_kb,
            elapsed * 1000,
            False,
        )
        with cache_lock:
            if memory_cache:
                memory_cache.set(cache_key, payload, payload_bytes)
        maybe_flush_query_stats()

        if is_fixed_candidate(query_stats.get(cache_key, {})) and cache_key not in fixed_cache_index:
            trigger_fixed_cache_rebuild(loop, "クエリ条件到達")

        # UI側でハイライトしやすいように検索キーワードも返す
        return payload


@app.post("/api/export")
async def export_results(req: SearchRequest, format: str = "csv"):
    """検索結果をCSV/JSON形式でエクスポートする"""
    global rw_lock, search_semaphore, search_executor, search_execution_mode
    if rw_lock is None or search_semaphore is None:
        raise HTTPException(status_code=503, detail="検索システムの初期化中です")

    # Validate format
    if format not in ["csv", "json"]:
        raise HTTPException(status_code=400, detail="format は csv または json を指定してください")

    # Execute search first
    async with search_semaphore:
        async with rw_lock.read_lock():
            available_ids = set(folder_states.keys())
            target_ids = [f for f in req.folders if f in available_ids and folder_states[f].get("ready")]
            if not target_ids:
                raise HTTPException(status_code=400, detail="有効な検索対象フォルダがありません")

            normalize_mode = resolve_normalize_mode(req.normalize_mode)
            params = SearchParams(req.mode, req.range_limit, req.space_mode, normalize_mode)
            keywords, norm_keyword_groups = build_query_groups(
                req.query,
                req.space_mode,
                normalize_mode,
            )
            raw_keywords = keywords
            worker_count = per_request_workers()
            loop = asyncio.get_running_loop()

            if search_execution_mode == "process":
                results = await loop.run_in_executor(
                    search_executor,
                    perform_search_process,
                    target_ids,
                    params,
                    norm_keyword_groups,
                    raw_keywords,
                    worker_count,
                    True,
                )
            else:
                results = await loop.run_in_executor(
                    None,
                    perform_search,
                    target_ids,
                    params,
                    norm_keyword_groups,
                    raw_keywords,
                    worker_count,
                    True,
                )

    # Get current index UUID
    current_gen_name = get_current_generation_pointer()
    index_uuid = current_gen_name or "unknown"

    # Prepare metadata
    metadata = {
        "query": req.query,
        "mode": req.mode,
        "range_limit": req.range_limit,
        "space_mode": req.space_mode,
        "normalize_mode": normalize_mode,
        "folders": target_ids,
        "folder_names": [folder_states[fid]["name"] for fid in target_ids if fid in folder_states],
        "index_uuid": index_uuid,
        "result_count": len(results),
        "exported_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }

    if format == "csv":
        # Create CSV
        output = io.StringIO()
        writer = csv.writer(output)

        # Write metadata as comments
        writer.writerow(["# フォルダ内テキスト検索 - エクスポート"])
        writer.writerow(["# クエリ", req.query])
        writer.writerow(["# モード", req.mode])
        writer.writerow(["# 範囲", req.range_limit])
        writer.writerow(["# 空白除去", req.space_mode])
        writer.writerow(["# 表記ゆれ", normalize_mode_label(normalize_mode)])
        writer.writerow(["# インデックスUUID", index_uuid])
        writer.writerow(["# エクスポート日時", metadata["exported_at"]])
        writer.writerow(["# 検索フォルダ", ", ".join(metadata["folder_names"])])
        writer.writerow(["# 結果件数", len(results)])
        writer.writerow([])

        # Write header
        writer.writerow(["ファイル名", "パス", "ページ", "フォルダ", "スニペット", "詳細"])

        # Write results
        for result in results:
            writer.writerow([
                result.get("file", ""),
                result.get("displayPath") or result.get("path", ""),
                result.get("page", ""),
                result.get("folderName", ""),
                result.get("context", ""),
                result.get("detail", ""),
            ])

        csv_content = output.getvalue()
        output.close()

        # Return as downloadable file
        return StreamingResponse(
            io.BytesIO(csv_content.encode("utf-8-sig")),  # UTF-8 with BOM for Excel compatibility
            media_type="text/csv; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="search_results_{time.strftime("%Y%m%d_%H%M%S")}.csv"'
            }
        )

    else:  # json
        # Create JSON with metadata
        export_data = {
            "metadata": metadata,
            "results": results,
        }

        json_content = json.dumps(export_data, ensure_ascii=False, indent=2)

        return StreamingResponse(
            io.BytesIO(json_content.encode("utf-8")),
            media_type="application/json; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="search_results_{time.strftime("%Y%m%d_%H%M%S")}.json"'
            }
        )


@app.get("/api/health")
async def health():
    ready = all(m.get("ready") for m in folder_states.values()) if folder_states else False
    return {"status": "ok", "ready": ready}


# --- 実行ヘルパー ---
def run():
    """ローカル開発用エントリポイント."""
    import uvicorn
    import signal

    port = int(os.getenv("PORT", "80"))
    cert_dir = os.getenv("CERT_DIR", "certs")
    base_dir = Path(__file__).resolve().parent
    cert_path = (base_dir / cert_dir).resolve()
    cert_file = cert_path / "lan-cert.pem"
    key_file = cert_path / "lan-key.pem"

    ssl_kwargs = {}
    if cert_file.exists() and key_file.exists():
        ssl_kwargs = {
            "ssl_certfile": str(cert_file),
            "ssl_keyfile": str(key_file),
        }

    config = uvicorn.Config(
        "web_server:app",
        host="0.0.0.0",
        port=port,
        reload=False,
        **ssl_kwargs,
    )
    server = uvicorn.Server(config)
    server.install_signal_handlers = lambda: None
    interrupt_count = {"count": 0}

    def handle_interrupt(signum, frame):
        interrupt_count["count"] += 1
        if interrupt_count["count"] == 1:
            log_info("Ctrl+Cを検知しました。もう一度Ctrl+Cで強制終了します。")
            server.should_exit = True
            return
        log_info("Ctrl+Cを再度検知しました。強制終了します。")
        os._exit(1)

    signal.signal(signal.SIGINT, handle_interrupt)
    try:
        signal.signal(signal.SIGTERM, handle_interrupt)
    except Exception:
        pass

    server.run()


if __name__ == "__main__":
    run()
