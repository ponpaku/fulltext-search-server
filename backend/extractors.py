"""File reading and text extraction utilities."""
from __future__ import annotations

import csv
import os
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from openpyxl import load_workbook
from pdfminer.high_level import extract_pages, extract_text as miner_extract_text
from pdfminer.layout import LAParams, LTTextBoxVertical, LTTextContainer
from pdfminer.pdfpage import PDFPage

from .text_utils import split_non_pdf_text

SYSTEM_VERSION = "1.1.11"
# File Version: 1.0.0


def read_text_file_safe(path: str) -> str:
    encodings = ["utf-8", "cp932", "shift_jis", "euc-jp"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""


def read_csv_file_safe(path: str) -> str:
    encodings = ["utf-8", "cp932", "shift_jis"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc, newline="") as f:
                reader = csv.reader(f)
                return "\n".join([" ".join(row) for row in reader])
        except Exception:
            continue
    return ""


def read_excel_rows_xlsx(path: str) -> List[Tuple[str, int, str]]:
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return []

    rows: List[Tuple[str, int, str]] = []
    try:
        for ws in wb.worksheets:
            sheet_name = ws.title or "Sheet"
            row_idx = 0
            for row in ws.iter_rows(values_only=True):
                row_idx += 1
                cells = [
                    str(cell).strip()
                    for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if cells:
                    rows.append((sheet_name, row_idx, " ".join(cells)))
    except Exception:
        return []
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return rows


def read_excel_rows_xls(path: str) -> List[Tuple[str, int, str]]:
    try:
        import xlrd
    except Exception:
        return []
    try:
        wb = xlrd.open_workbook(path)
    except Exception:
        return []

    rows: List[Tuple[str, int, str]] = []
    try:
        for sheet in wb.sheets():
            sheet_name = sheet.name or "Sheet"
            for r in range(sheet.nrows):
                row_values = [
                    str(sheet.cell_value(r, c)).strip()
                    for c in range(sheet.ncols)
                    if str(sheet.cell_value(r, c)).strip()
                ]
                if row_values:
                    rows.append((sheet_name, r + 1, " ".join(row_values)))
    except Exception:
        return []
    return rows


def split_excel_rows(rows: List[Tuple[str, int, str]], rows_per_page: int = 40) -> Dict[str, str]:
    pages: Dict[str, str] = {}
    by_sheet: Dict[str, List[Tuple[int, str]]] = {}
    for sheet, row_idx, text in rows:
        by_sheet.setdefault(sheet, []).append((row_idx, text))

    for sheet, sheet_rows in by_sheet.items():
        sheet_rows.sort(key=lambda x: x[0])
        for i in range(0, len(sheet_rows), rows_per_page):
            chunk = sheet_rows[i : i + rows_per_page]
            start_row = chunk[0][0]
            end_row = chunk[-1][0]
            label = f"{sheet} 行 {start_row}-{end_row}"
            pages[label] = "\n".join(text for _, text in chunk).strip()
    return pages


def iter_text_containers(obj):
    if isinstance(obj, LTTextContainer):
        yield obj
        return
    for child in getattr(obj, "_objs", []):
        yield from iter_text_containers(child)


def order_text_boxes(boxes: List[Dict], page_bbox: Tuple[float, float, float, float]) -> str:
    if not boxes:
        return ""
    x0, y0, x1, y1 = page_bbox
    page_mid = (x0 + x1) / 2.0
    vertical_count = sum(1 for b in boxes if b["vertical"])
    horizontal_count = len(boxes) - vertical_count
    use_vertical = vertical_count > horizontal_count and vertical_count >= 3

    if use_vertical:
        ordered = sorted(boxes, key=lambda b: (-b["x0"], -b["y1"]))
    else:
        left = [b for b in boxes if b["cx"] < page_mid]
        right = [b for b in boxes if b["cx"] >= page_mid]
        is_two_column = (
            len(boxes) >= 6
            and len(left) >= 2
            and len(right) >= 2
            and 0.3 < (len(left) / len(boxes)) < 0.7
        )
        if is_two_column:
            left.sort(key=lambda b: (-b["y1"], b["x0"]))
            right.sort(key=lambda b: (-b["y1"], b["x0"]))
            ordered = left + right
        else:
            ordered = sorted(boxes, key=lambda b: (-b["y1"], b["x0"]))

    return "\n".join(b["text"].strip() for b in ordered if b["text"].strip())


def extract_pdf_text_by_layout(file_path: str) -> Dict[int, str]:
    page_texts: Dict[int, str] = {}
    laparams = LAParams()
    for page_num, layout in enumerate(extract_pages(file_path, laparams=laparams), start=1):
        boxes = []
        for container in iter_text_containers(layout):
            text = container.get_text()
            if not text or not text.strip():
                continue
            x0, y0, x1, y1 = container.bbox
            boxes.append(
                {
                    "text": text,
                    "x0": x0,
                    "y0": y0,
                    "x1": x1,
                    "y1": y1,
                    "cx": (x0 + x1) / 2.0,
                    "vertical": isinstance(container, LTTextBoxVertical),
                }
            )
        if boxes:
            ordered_text = order_text_boxes(boxes, layout.bbox)
            if ordered_text:
                page_texts[page_num] = ordered_text
    return page_texts


def extract_text_from_file(file_path: str) -> Dict[int, str]:
    """Extract text from file (PDF uses high-accuracy mode)."""
    texts, _ = extract_text_from_file_with_reason(file_path)
    return texts


def empty_reason(file_path: str, fallback: str) -> str:
    try:
        if os.path.getsize(file_path) == 0:
            return "空ファイル"
    except Exception:
        return fallback
    return fallback


def extract_text_from_file_with_reason(file_path: str) -> Tuple[Dict[int, str], str]:
    """Extract text and return reason if failed."""
    path_obj = Path(file_path)
    ext = path_obj.suffix.lower()
    page_texts: Dict[int, str] = {}
    file_path_str = str(file_path)
    reason = ""

    try:
        if ext == ".pdf":
            try:
                page_texts = extract_pdf_text_by_layout(file_path_str)
            except Exception:
                page_texts = {}
            if not page_texts:
                laparams = LAParams()
                with open(file_path_str, "rb") as fp:
                    pages = list(PDFPage.get_pages(fp, check_extractable=False))
                    total_pages = len(pages)

                with open(file_path_str, "rb") as fp:
                    for i in range(total_pages):
                        try:
                            text = miner_extract_text(fp, page_numbers=[i], laparams=laparams)
                            if text and text.strip():
                                page_texts[i + 1] = text
                        except Exception:
                            continue
            if not page_texts:
                reason = empty_reason(file_path_str, "PDF抽出失敗")

        elif ext == ".docx":
            doc = Document(file_path_str)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            content = "\n".join(full_text)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "内容が空")

        elif ext == ".txt":
            content = read_text_file_safe(file_path_str)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")

        elif ext == ".csv":
            content = read_csv_file_safe(file_path_str)
            if content:
                page_texts = split_non_pdf_text(content)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
        elif ext == ".xlsx":
            rows = read_excel_rows_xlsx(file_path_str)
            if rows:
                page_texts = split_excel_rows(rows)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
        elif ext == ".xls":
            rows = read_excel_rows_xls(file_path_str)
            if rows:
                page_texts = split_excel_rows(rows)
            else:
                reason = empty_reason(file_path_str, "読み取り失敗または内容が空")
    except Exception as exc:
        return {}, f"例外:{type(exc).__name__}"

    if page_texts:
        return page_texts, ""
    return {}, reason or "抽出結果が空"
